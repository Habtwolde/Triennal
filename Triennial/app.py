import os
import re
import json
import time
import base64
import hashlib
import zipfile
import tempfile
from datetime import datetime
import subprocess
from pathlib import Path
import collections
import shutil
import requests
from typing import Optional, Tuple
from functools import lru_cache
import xml.etree.ElementTree as ET
import calendar

import pandas as pd
import streamlit as st

# Optional DOCX formatting support (python-docx may not exist in Databricks Apps)
DOCX_AVAILABLE = True
try:
    from docx import Document
    from docx.shared import RGBColor
except Exception:
    DOCX_AVAILABLE = False


# =============================
# 0) App config
# =============================
st.set_page_config(page_title="Triennial Report Generator", layout="wide")
# =============================
# UI layout: constrain content width (prevents widgets from stretching to far right)
# =============================
st.markdown(
    """
    <style>
      /* Limit the content width even when layout="wide" */
      section.main > div.block-container{
        max-width: 1180px;     /* adjust: 980–1400 depending on preference */
        padding-left: 2.2rem;
        padding-right: 2.2rem;
      }

      /* Keep tables from looking edge-to-edge */
      div[data-testid="stDataFrame"]{
        max-width: 100%;
      }

      /* Optional: slightly reduce selectbox width feel */
      div[data-testid="stSelectbox"]{
        max-width: 100%;
      }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("Triennial Report Generator")
st.caption("Select inputs, then filter by Field to generate a publication-ready DOCX report.")

# =============================
# UI behavior
# =============================
NARRATE_EVERY_N_DEFAULT = 3
SHOW_PARTIAL_OUTPUT = False
SILENT_STAGING = True

# =============================
# LLM Token Limits
# =============================

# ROW_MAX_TOKENS = 800          # per-UID paragraph generation
# SUMMARY_MAX_TOKENS = 500      # summary generation (if used)
# PLAN_MAX_TOKENS = 1200        # plan generation (if used)

# =============================
# LLM Generation Controls
# =============================

ROW_MAX_TOKENS = 800
ROW_TEMP = 0.2

PLAN_MAX_TOKENS = 1200
PLAN_TEMP = 0.1

SUMMARY_MAX_TOKENS = 500
SUMMARY_TEMP = 0.1


# =============================
# Narrative shaping constants
# =============================
INTRO_MIN_PARAS = 2
INTRO_TARGET_MAX = 3
INTRO_MIN_WORDS = 170
INTRO_RETRY_LIMIT = 4

SUMMARY_MIN_PARAS = 2
SUMMARY_TARGET_MAX = 2
SUMMARY_MIN_WORDS = 140
SUMMARY_RETRY_LIMIT = 4


# =============================
# LLM endpoint & generation defaults
# =============================
# Databricks Model Serving endpoint name (override via env var TRIENNIAL_ENDPOINT).
ENDPOINT = (os.environ.get("TRIENNIAL_ENDPOINT") or "databricks-meta-llama-3-3-70b-instruct").strip()

# Default generation params (kept consistent with the original app.py)
TEMPERATURE = float(os.environ.get('TRIENNIAL_TEMPERATURE', '0.25'))
MAX_TOKENS_ROW = int(os.environ.get('TRIENNIAL_MAX_TOKENS_ROW', '550'))
MAX_TOKENS_SYN = int(os.environ.get('TRIENNIAL_MAX_TOKENS_SYN', '450'))
MAX_TOKENS_INTRO = int(os.environ.get('TRIENNIAL_MAX_TOKENS_INTRO', '1200'))
MAX_TOKENS_SUMMARY = int(os.environ.get('TRIENNIAL_MAX_TOKENS_SUMMARY', '900'))

# Sentence-ending punctuation detector used by post-processing
_END_PUNCT_RE = re.compile(r"[.!?][\'\"\)\]]?\s*$")
# Canonical column names expected from the master Excel file
CANON = [
    "Submitting ICO", "Lead ICO", "Unique ID", "Collaborating ICOs/Agencies/Orgs",
    "Activity Name", "Activity Description", "Activity Type", "Field", "Importance",
    "Web address(es)", "PMID(s)", "Notes", "Notes.1"
]

# Field ordering used for routing & report section sequence
SECTION_ORDER = [
    "Advanced Imaging & AI Tools",
    "Combination & Targeted Therapies",
    "Data Commons and Computational Resources",
    "Environmental Health and Cancer",
    "Epidemiology & Surveillance",
    "Genetics, Cell Biology, and -Omics",
    "Immunotherapy",
    "Nutrition & Symptom Management",
    "Preventive Interventions",
    "Recalcitrant & Hard-to-Treat Cancer Research",
    "Screening & Early Detection",
    "Tumor Microenvironment & Immunology",
]

# =============================
# Acronym handling
# =============================
DISABLE_ACRONYM_EXPANSION = True  # Client requirement: do NOT expand acronyms in narrative text
PROTECTED_ACRONYMS = {"CI", "CIs", "C.I.", "C.I.s", "C.I.s.", "C.I", "CI.", "CIs."}

NO_ACRONYM_EXPANSION_INSTRUCTION = (
    "Do NOT expand acronyms in the narrative text. Keep acronyms exactly as written in the source. "
    "Never expand CI or CIs."
)
# =============================
# 1) Default paths & constants
# =============================
# Client portability: default to reading inputs from the *app folder* (the folder that contains this app.py).
# This avoids hard-coding DBFS paths that differ across workspaces.
APP_DIR = Path(__file__).resolve().parent
APP_ASSETS_DIR = (APP_DIR / "assets") if (APP_DIR / "assets").exists() else APP_DIR

import shutil
print("Pandoc on PATH:", shutil.which("pandoc"))

def _app_path(p: str) -> str:
    """Resolve a user-supplied path. Relative paths are resolved under APP_ASSETS_DIR."""
    p = (p or "").strip()
    if not p:
        return ""
    if p.startswith("dbfs:") or p.startswith("/dbfs/"):
        return p  # leave DBFS paths untouched (optional support)
    if p.startswith("/"):
        return p
    return str(APP_ASSETS_DIR / p)

# Defaults expect these files to be shipped alongside app.py (or under ./assets/)
DEFAULT_EXCEL_PATH = _app_path("Triennial Data Source_Master File of All Submissions_OEPR Ch3 Writers (1).xlsx")
DEFAULT_STYLE_PROMPT_PATH = _app_path("style_prompt.json")
DEFAULT_REFERENCE_DOCX_PATH = _app_path("reference.docx")
DEFAULT_LUA_FILTER_PATH = _app_path("h2_pagebreak.lua")
DEFAULT_SQUARE_FILTER_PATH = _app_path("h2_square_bracket_footnotes.lua")

# Output options (DBFS publish is optional; download_button works without DBFS)
DEFAULT_WORKING_OUT_DBFS = "dbfs:/FileStore/triennial/out"
DEFAULT_VOLUME_OUT_DIR = "/Volumes/dpcpsi/gold/triennial_reports"

# Local, writable staging area inside the App container
LOCAL_ASSETS_DIR = "/tmp/triennial_assets"
LOCAL_OUT_DIR = "/tmp/triennial_out"
Path(LOCAL_ASSETS_DIR).mkdir(parents=True, exist_ok=True)
Path(LOCAL_OUT_DIR).mkdir(parents=True, exist_ok=True)

EXCEL_LOCAL = str(Path(LOCAL_ASSETS_DIR) / "master.xlsx")
STYLE_PROMPT_LOCAL = str(Path(LOCAL_ASSETS_DIR) / "style_prompt.json")
REFERENCE_DOCX_LOCAL = str(Path(LOCAL_ASSETS_DIR) / "reference.docx")
LUA_FILTER_LOCAL = str(Path(LOCAL_ASSETS_DIR) / "h2_pagebreak.lua")
SQUARE_FILTER_LOCAL = str(Path(LOCAL_ASSETS_DIR) / "h2_square_bracket_footnotes.lua")


# =============================
# 2) Auth (OAuth from App environment)
# =============================
def _env(name: str) -> str:
    return (os.environ.get(name) or "").strip()

def get_workspace_host() -> str:
    host = _env("DATABRICKS_HOST")
    if host and not host.startswith("http"):
        host = "https://" + host
    return host.rstrip("/")

@st.cache_resource(show_spinner=False)
def get_oauth_token() -> str:
    host = get_workspace_host()
    client_id = _env("DATABRICKS_CLIENT_ID")
    client_secret = _env("DATABRICKS_CLIENT_SECRET")

    if not host or not client_id or not client_secret:
        raise RuntimeError(
            "Missing OAuth env vars inside the App container.\n"
            "Required: DATABRICKS_HOST, DATABRICKS_CLIENT_ID, DATABRICKS_CLIENT_SECRET."
        )

    token_url = f"{host}/oidc/v1/token"
    data = {"grant_type": "client_credentials", "scope": "all-apis"}

    r = requests.post(token_url, data=data, auth=(client_id, client_secret), timeout=30)
    r.raise_for_status()
    return r.json()["access_token"]

def auth_headers() -> dict:
    return {"Authorization": f"Bearer {get_oauth_token()}"}

# =============================
# 3) DBFS REST helpers (Apps-safe)
# =============================
def dbfs_norm(dbfs_path: str) -> str:
    return dbfs_path

def dbfs_get_status(dbfs_path: str) -> dict:
    host = get_workspace_host()
    url = f"{host}/api/2.0/dbfs/get-status"
    r = requests.get(url, headers=auth_headers(), params={"path": dbfs_norm(dbfs_path)}, timeout=30)
    r.raise_for_status()
    return r.json()

def dbfs_read_all(dbfs_path: str, chunk_size: int = 1_000_000) -> bytes:
    host = get_workspace_host()
    url = f"{host}/api/2.0/dbfs/read"
    offset = 0
    out = bytearray()

    while True:
        r = requests.get(
            url,
            headers=auth_headers(),
            params={"path": dbfs_norm(dbfs_path), "offset": offset, "length": chunk_size},
            timeout=60,
        )
        r.raise_for_status()
        j = r.json()

        data_b64 = j.get("data", "") or ""
        data = base64.b64decode(data_b64) if data_b64 else b""
        out.extend(data)

        bytes_read = j.get("bytes_read", 0) or 0
        if bytes_read <= 0:
            break
        offset += bytes_read
        if j.get("eof", False):
            break

    return bytes(out)

def dbfs_write_file(local_path: str, content: bytes):
    p = Path(local_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_bytes(content)

def dbfs_mkdirs(dbfs_dir: str):
    host = get_workspace_host()
    url = f"{host}/api/2.0/dbfs/mkdirs"
    r = requests.post(url, headers=auth_headers(), json={"path": dbfs_norm(dbfs_dir)}, timeout=30)
    r.raise_for_status()

def dbfs_delete_if_exists(dbfs_path: str):
    host = get_workspace_host()
    url = f"{host}/api/2.0/dbfs/delete"
    r = requests.post(
        url,
        headers=auth_headers(),
        json={"path": dbfs_norm(dbfs_path), "recursive": False},
        timeout=30,
    )
    if r.status_code not in (200, 404):
        r.raise_for_status()

def dbfs_create(dbfs_path: str, overwrite: bool = True) -> int:
    host = get_workspace_host()
    url = f"{host}/api/2.0/dbfs/create"

    if overwrite:
        dbfs_delete_if_exists(dbfs_path)

    r = requests.post(url, headers=auth_headers(), json={"path": dbfs_norm(dbfs_path)}, timeout=30)
    r.raise_for_status()
    return int(r.json()["handle"])

def dbfs_add_block(handle: int, data_block: bytes):
    host = get_workspace_host()
    url = f"{host}/api/2.0/dbfs/add-block"
    payload = {"handle": handle, "data": base64.b64encode(data_block).decode("utf-8")}
    r = requests.post(url, headers=auth_headers(), json=payload, timeout=60)
    r.raise_for_status()

def dbfs_close(handle: int):
    host = get_workspace_host()
    url = f"{host}/api/2.0/dbfs/close"
    r = requests.post(url, headers=auth_headers(), json={"handle": handle}, timeout=30)
    r.raise_for_status()

def dbfs_put_large(dbfs_path: str, data: bytes, overwrite: bool = True, block_size: int = 1_000_000):
    handle = dbfs_create(dbfs_path, overwrite=overwrite)
    try:
        for i in range(0, len(data), block_size):
            dbfs_add_block(handle, data[i:i + block_size])
    finally:
        dbfs_close(handle)

def stage_assets_or_stop(
    excel_path: str,
    style_prompt_path: str,
    reference_docx_path: str,
    lua_filter_path: str,
    square_filter_path: str,
):
    """
    Stage required assets into the App container's writable area (/tmp).

    - Preferred mode: read from local app folder (relative paths under APP_ASSETS_DIR).
    - Optional compatibility: if a path starts with dbfs:/ or /dbfs/, read via DBFS REST.
    """

    def _is_dbfs(p: str) -> bool:
        p = (p or "").strip()
        return p.startswith("dbfs:") or p.startswith("/dbfs/")

    def _read_src_bytes(src: str) -> bytes:
        src = (src or "").strip()
        if not src:
            return b""
        if _is_dbfs(src):
            return dbfs_read_all(src)
        # local file
        return Path(src).read_bytes()

    def _stage_one(src: str, dst_local: str, required: bool = True, label: str = "") -> Tuple[bool, Optional[str], Optional[str]]:
        """Returns (ok, error_kind, error_detail)."""
        src = (src or "").strip()
        if not src:
            if required:
                return False, "missing", f"{label or 'File'} path is empty."
            return True, None, None
        try:
            if _is_dbfs(src):
                status = dbfs_get_status(src)
                data = _read_src_bytes(src)
                if not data:
                    raise RuntimeError("Empty read from DBFS.")
                sha = hashlib.sha256(data).hexdigest()
                dbfs_write_file(dst_local, data)
                staged_info.append((src, dst_local, status.get("file_size"), status.get("modification_time"), sha))
                return True, None, None
            else:
                p = Path(src)
                if not p.exists():
                    raise FileNotFoundError(f"Local path does not exist: {p}")
                data = _read_src_bytes(src)
                if not data:
                    raise RuntimeError("Empty read from local file.")
                sha = hashlib.sha256(data).hexdigest()
                Path(dst_local).parent.mkdir(parents=True, exist_ok=True)
                Path(dst_local).write_bytes(data)
                staged_info.append((str(p), dst_local, len(data), int(p.stat().st_mtime * 1000), sha))
                return True, None, None
        except Exception as e:
            if required:
                return False, "unreadable", f"{label or 'File'} could not be read: {src}\n{e}"
            return True, None, None

    staged_info = []  # (src, dst_local, file_size, mtime, sha256)

    # Resolve app-relative paths (so users can type "style_prompt.json" etc.)
    excel_path = _app_path(excel_path)
    style_prompt_path = _app_path(style_prompt_path)
    reference_docx_path = _app_path(reference_docx_path)
    lua_filter_path = _app_path(lua_filter_path)
    square_filter_path = _app_path(square_filter_path)

    missing_or_bad = []

    ok, kind, detail = _stage_one(excel_path, EXCEL_LOCAL, required=True, label="Excel")
    if not ok: missing_or_bad.append(detail)

    ok, kind, detail = _stage_one(style_prompt_path, STYLE_PROMPT_LOCAL, required=True, label="Style prompt JSON")
    if not ok: missing_or_bad.append(detail)

    ok, kind, detail = _stage_one(reference_docx_path, REFERENCE_DOCX_LOCAL, required=True, label="Reference DOCX")
    if not ok: missing_or_bad.append(detail)

    # Lua filters are optional; pagebreak is used in the current export path, square-bracket filter is staged for future use.
    _stage_one(lua_filter_path, LUA_FILTER_LOCAL, required=False, label="Lua pagebreak filter")
    _stage_one(square_filter_path, SQUARE_FILTER_LOCAL, required=False, label="Lua square-bracket filter")

    if missing_or_bad:
        st.error(
            "One or more REQUIRED input files could not be staged into the App container.\n\n"
            + "\n\n".join(missing_or_bad)
            + "\n\n"
            "Tip: Put the files next to app.py (or under ./assets/) and use relative paths like 'style_prompt.json'."
        )
        st.stop()

    if not SILENT_STAGING:
        st.success("Assets staged into App container (/tmp).")

        with st.expander("Staged asset verification (source → /tmp)", expanded=False):
            if not staged_info:
                st.write("No asset metadata captured.")
            else:
                rows = []
                for src, dst, sz, mt, sha in staged_info:
                    rows.append(
                        {
                            "Source": src,
                            "Staged to": dst,
                            "Bytes": sz,
                            "Modified (ms)": mt,
                            "SHA-256": sha,
                        }
                    )
                st.dataframe(pd.DataFrame(rows), use_container_width=True)

def _trim_to_last_complete_sentence(paragraph: str) -> str:
    """
    If the paragraph looks cut off (no terminal punctuation), trim back to the
    last sentence boundary. This is deterministic and prevents abrupt endings
    in DOCX when the model truncates due to token limits.
    """
    if not paragraph:
        return paragraph

    p = re.sub(r"\s+", " ", paragraph).strip()
    if not p:
        return p

    # Already ends well
    if _END_PUNCT_RE.search(p):
        return p

    # Try to trim to the last sentence-ending punctuation.
    last = max(p.rfind("."), p.rfind("!"), p.rfind("?"))
    if last >= 0 and last >= int(len(p) * 0.60):
        p = p[: last + 1].strip()
        if _END_PUNCT_RE.search(p):
            return p

    # Fallback: add a period
    return p.rstrip(";,:-") + "."

def finalize_multparagraph_text(md_text: str) -> str:
    """
    Apply clean endings paragraph-by-paragraph.
    """
    if not md_text:
        return md_text

    paras = [p.strip() for p in re.split(r"\n\s*\n", md_text) if p.strip()]
    cleaned = [_trim_to_last_complete_sentence(p) for p in paras]
    return "\n\n".join(cleaned).strip()

def _safe_filename(s: str) -> str:
    s = (s or "output").strip()
    s = re.sub(r"[^A-Za-z0-9._-]+", "_", s)
    return s.strip("_")[:80] or "output"

def build_uid_entry_map(section_order: list[str], section_to_uids: dict) -> dict[str, int]:
    """
    Deterministic global entry numbering across the report, in the order
    sections are printed and UIDs appear within each section.
    """
    uid_to_entry: dict[str, int] = {}
    n = 0
    for sec in section_order:
        for uid in (section_to_uids.get(sec, []) or []):
            if uid in uid_to_entry:
                continue
            n += 1
            uid_to_entry[uid] = n
    return uid_to_entry


def normalize_uid_marker_placement(text: str) -> str:
    """
    Conservative normalization of UID marker placement.
    Fixes common pattern: 'word[^UID].' -> 'word.[^UID]'
    Does NOT attempt full sentence parsing.
    """
    if not text:
        return text

    # Move marker that appears before punctuation to after it.
    text = re.sub(r"(\[\^\s*[A-Za-z0-9._-]+\s*\])([.,;:])", r"\2\1", text)
    return text

def add_uid_marker_per_sentence(text: str, uid: str) -> str:
    """
    Place the UID footnote marker at the end of EACH sentence (not just paragraph end).

    Rules:
      - Strip any existing [^...] markers first (we re-insert deterministically).
      - Protect common abbreviations and initialisms so we don't treat their periods as sentence boundaries.
      - Insert [^<uid>] immediately after sentence-ending punctuation (.!?).
    """
    import re

    if not text:
        return text

    # Remove any existing footnote markers first
    s = re.sub(r"\[\^\s*[A-Za-z0-9._-]+\s*\]", "", text).strip()

    # Token used to protect periods that are NOT sentence boundaries
    DOT = "∯"

    # Protect common abbreviations (extend as needed)
    abbrevs = [
        "e.g.", "i.e.", "etc.", "vs.", "Dr.", "Mr.", "Ms.", "Mrs.", "Prof.", "Sr.", "Jr.",
        "Fig.", "No.", "St.", "a.m.", "p.m."
    ]
    for a in abbrevs:
        s = s.replace(a, a.replace(".", DOT))

    # Protect initialisms like U.S., U.K., N.I.H., etc.
    def _protect_initialisms(m):
        return m.group(0).replace(".", DOT)
    s = re.sub(r"\b(?:[A-Z]\.){2,}", _protect_initialisms, s)

    # Insert marker after sentence-ending punctuation.
    # Handles optional closing quotes/brackets before whitespace/end.
    def _ins(m):
        punct = m.group(1)
        tail = m.group(2) or ""
        return f"{punct}[^{uid}]{tail}"

    s = re.sub(r"([.!?])([\"'\)\]\}]*)(\s+|$)", lambda m: _ins(m) if m.group(3) is not None else m.group(0), s)

    # Restore protected periods
    s = s.replace(DOT, ".")

    return s


def apply_uid_markers_from_cite_tokens(text: str, uid: str, token: str = "[[CITE]]") -> str:
    """Apply UID markers ONLY where the model indicated direct support.

    The model is instructed to append `[[CITE]]` to sentences that are directly supported
    by the activity brief. We convert those tokens into the UID footnote marker.

    Behavior:
      - Strip any existing [^...] markers first (deterministic reinsertion).
      - Replace `[[CITE]]` tokens with `[^<uid>]`, preserving punctuation order.
      - If NO token appears, leave the paragraph uncited (no paragraph-end fallback).
    """
    import re

    if not text:
        return text

    s = re.sub(r"\[\^\s*[A-Za-z0-9._-]+\s*\]", "", text).strip()

    # Replace token patterns near punctuation/whitespace
    # e.g., "sentence.[[CITE]]" or "sentence. [[CITE]]"
    token_re = re.escape(token)

    # Count tokens before replacement (audit)
    n_tokens = len(re.findall(token_re, s))

    # Normalize ".[[CITE]]" -> ".[^UID]" and " [[CITE]]" -> "[^UID]"
    s = re.sub(rf"([.!?])\s*{token_re}", rf"\1[^{uid}]", s)
    s = re.sub(rf"\s*{token_re}", rf"[^{uid}]", s)

    # Remove any leftover token text (if model altered it slightly)
    s = s.replace(token, "").strip()

    if n_tokens == 0:
        # Client requirement: do NOT add paragraph-end citations when the model did not mark a sentence as supported.
        return s

    return s

def strip_raw_uid_tokens(text: str) -> str:
    """
    Remove raw UID tokens (e.g., '378_NIAID', '697_NHLBI') from narrative prose.

    Critical: DO NOT strip the UID inside Pandoc footnote markers like [^378_NIAID].
    If we remove the UID from inside the marker, the text becomes '[^]' and Pandoc
    will no longer convert it into a footnote/citation (it will remain literal in the DOCX).
    """
    if not text:
        return text

    # 1) Protect existing footnote markers
    markers: list[str] = []

    def _protect(m: re.Match) -> str:
        markers.append(m.group(0))
        return f"__UIDMARKER_{len(markers)-1}__"

    protected = re.sub(r"\[\^\s*[A-Za-z0-9._-]+\s*\]", _protect, text)

    # 2) Strip standalone UID tokens (not inside markers)
    protected = re.sub(r"\b\d+_[A-Za-z0-9]+\b", "", protected)

    # 3) Clean spacing artifacts
    protected = re.sub(r"\s{2,}", " ", protected)
    protected = re.sub(r"\s+([\.,;:!?])", r"\1", protected)

    # 4) Restore markers
    def _restore(m: re.Match) -> str:
        idx = int(m.group(1))
        return markers[idx] if 0 <= idx < len(markers) else m.group(0)

    restored = re.sub(r"__UIDMARKER_(\d+)__", _restore, protected)
    return restored.strip()


def inject_entry_labels_near_uid_markers(text: str, uid_to_entry: dict[str, int]) -> str:
    """
    Replace each UID marker [^UID] with '(Entry n)[^UID]' so reviewers can trace each citation
    back to a specific activity entry. If a marker is already preceded by an Entry label
    within a short window, it is left unchanged (prevents double-insertion).
    """
    if not text:
        return text

    UID_MARK_RE = re.compile(r"\[\^\s*([A-Za-z0-9._-]+)\s*\]")

    def _repl(m: re.Match) -> str:
        uid = m.group(1)
        n = uid_to_entry.get(uid)
        if not n:
            return m.group(0)

        # Prevent double insertion if an Entry label is already close to the marker
        start = max(0, m.start() - 25)
        window = text[start:m.start()]
        if re.search(r"\(Entry\s+\d+\)\s*$", window):
            return m.group(0)

        return f"(Entry {n})" + m.group(0)

    return UID_MARK_RE.sub(_repl, text)

def nih_style_citation_phrasing(text: str) -> str:
    """
    Strict NIH narrative citation style:
      - No lead-in/meta prose that explains citations.
      - Multiple citations separated by COMMAS.
    """
    if not text:
        return text

    text = re.sub(r"\s+", " ", text).strip()

    leadins = [
        r"as\s+evidenced\s+by",
        r"evidenced\s+by",
        r"marked\s+by",
        r"marked",
        r"illustrated\s+by",
        r"as\s+illustrated\s+by",
        r"demonstrated\s+by",
        r"as\s+demonstrated\s+by",
        r"supported\s+by",
        r"as\s+supported\s+by",
        r"informed\s+by",
        r"as\s+informed\s+by",
        r"including\s+via",
        r"including\s+through",
        r"including",
        r"such\s+as",
        r"for\s+example",
        r"e\.g\.",
        r"like",
        r"as\s+seen\s+in",
        r"as\s+seen\s+through",
        r"as\s+seen\s+via",
        r"as\s+seen",
        r"as\s+indicated\s+by",
        r"as\s+indicated",
        r"as\s+shown\s+by",
        r"as\s+shown",
        r"activities\s+such\s+as\s+those\s+associated\s+with",
        r"those\s+associated\s+with",
        r"associated\s+with",
        r"as\s+indicated\s+by\s+the\s+\w+\s+unique\s+uids",
        r"unique\s+uids",
    ]
    for li in leadins:
        text = re.sub(rf"\b{li}\b\s*(?=\(Entry\s+\d+\)\[\^)", "", text, flags=re.IGNORECASE)
        text = re.sub(rf"\b{li}\b\s*(?=\[\^)", "", text, flags=re.IGNORECASE)

    text = re.sub(r"\b(and|or)\s*,?\s*(?=\(Entry\s+\d+\)\[\^)", "", text, flags=re.IGNORECASE)
    text = re.sub(r",\s*(?=\(Entry\s+\d+\)\[\^)", "", text)

    text = re.sub(r"\s*,\s*(?=\(Entry\s+\d+\)\[\^)", " ", text)
    text = re.sub(r"\.(?=\(Entry\s+\d+\)\[\^)", "", text)

    cite_token = r"\(Entry\s+\d+\)\[\^[A-Za-z0-9._-]+\]"
    seq_re = re.compile(rf"({cite_token})(?:\s*,?\s*({cite_token}))+")

    def _format_seq(m: re.Match) -> str:
        tokens = re.findall(cite_token, m.group(0))
        return ", ".join(tokens)

    text = seq_re.sub(_format_seq, text)

    text = re.sub(r"\s+([\.,;:!?])", r"\1", text)
    text = re.sub(r"\(\s+", "(", text)
    text = re.sub(r"\s+\)", ")", text)

    return text

def ensure_period_after_citations(text: str) -> str:
    """
    Ensure sentence-ending period appears AFTER citation blocks.
    """
    if not text:
        return text

    text = re.sub(
        r"\.\s*((?:\((?:Entry\s+\d+)\)\[\^[^\]]+\](?:,\s*)?)+)",
        r" \1.",
        text
    )

    text = re.sub(
        r"((?:\((?:Entry\s+\d+)\)\[\^[^\]]+\](?:,\s*)?)+)(?=\s*$)",
        r"\1.",
        text
    )

    return text


def ensure_semicolon_after_entry_citation(text: str) -> str:
    """
    Ensure a grammatical boundary after Entry citations when a new sentence/clause starts.
    """
    if not text:
        return text

    text = re.sub(
        r"(\(Entry\s+\d+\)\[\^[^\]]+\])\s+(?=[A-Z])",
        r"\1; ",
        text
    )
    text = re.sub(
        r"(\(Entry\s+\d+\))\s+(?=[A-Z])",
        r"\1; ",
        text
    )
    return text


def replace_percent_sign(text: str) -> str:
    """
    Convert percent signs to the word 'percent' per client requirement.
    """
    if not text:
        return text
    parts = re.split(r"(https?://\S+)", text)
    for i, p in enumerate(parts):
        if p.startswith("http://") or p.startswith("https://"):
            continue
        p = re.sub(r"(\d+(?:\.\d+)?)\s*%", r"\1 percent", p)
        if "%" in p:
            p = p.replace("%", " percent")
        parts[i] = p
    return "".join(parts)


_ICO_FULLNAME_TO_ACR = {
    "National Institutes of Health": "NIH",
    "National Cancer Institute": "NCI",
    "National Institute of Allergy and Infectious Diseases": "NIAID",
    "National Heart, Lung, and Blood Institute": "NHLBI",
    "National Center for Advancing Translational Sciences": "NCATS",
    "National Institute of Biomedical Imaging and Bioengineering": "NIBIB",
    "National Institute of Environmental Health Sciences": "NIEHS",
    "National Institute of Diabetes and Digestive and Kidney Diseases": "NIDDK",
    "National Institute on Aging": "NIA",
    "National Institute of Mental Health": "NIMH",
    "National Institute of Neurological Disorders and Stroke": "NINDS",
    "Eunice Kennedy Shriver National Institute of Child Health and Human Development": "NICHD",
    "National Institute on Alcohol Abuse and Alcoholism": "NIAAA",
    "National Institute of General Medical Sciences": "NIGMS",
    "National Human Genome Research Institute": "NHGRI",
    "National Library of Medicine": "NLM",
    "Office of the Director": "OD",
}


def revert_ci_expansions(text: str) -> str:
    """Revert accidental expansions of CI/CIs back to the acronyms in narrative text."""
    if not text:
        return text
    # Common variants; keep this intentionally narrow to avoid unintended edits.
    text = re.sub(r"\bconfidence intervals\b", "CIs", text, flags=re.IGNORECASE)
    text = re.sub(r"\bconfidence interval\b", "CI", text, flags=re.IGNORECASE)
    return text


def enforce_ico_acronyms(text: str) -> str:
    """
    Enforce client rule: NIH ICOs must appear as acronyms in narrative paragraphs.
    """
    if not text:
        return text

    for full, acr in _ICO_FULLNAME_TO_ACR.items():
        text = re.sub(
            rf"\b{re.escape(full)}\s*\(\s*{re.escape(acr)}\s*\)",
            acr,
            text,
        )
        text = re.sub(rf"\b{re.escape(full)}\b", acr, text)

    acr_pat = "|".join(sorted({v for v in _ICO_FULLNAME_TO_ACR.values()} | {"NIH"}))
    text = re.sub(rf"\bthe\s+({acr_pat})\b", r"\1", text, flags=re.IGNORECASE)

    return text


def postprocess_narrative(text: str) -> str:
    """
    Centralized deterministic cleanup applied AFTER LLM output and BEFORE DOCX rendering.
    """
    if not text:
        return text
    text = enforce_ico_acronyms(text)
    text = replace_percent_sign(text)
    text = ensure_semicolon_after_entry_citation(text)
    return text

def compact_consecutive_citations(t: str) -> str:
    """
    Make citation markers 'silent' and compact.
    """
    if not t:
        return t

    t = re.sub(r"\s+", " ", t).strip()
    t = re.sub(r"(\[\^\s*[A-Za-z0-9._-]+\s*\])\s*,\s*(\[\^\s*[A-Za-z0-9._-]+\s*\])", r"\1\2", t)
    t = re.sub(
        r"(\[\^\s*[A-Za-z0-9._-]+\s*\])\s*(?:,?\s*(?:and|or)\s+)\s*(\[\^\s*[A-Za-z0-9._-]+\s*\])",
        r"\1\2",
        t,
        flags=re.IGNORECASE
    )
    t = re.sub(r",\s*(?=\[\^)", " ", t)
    t = re.sub(r"\b(and|or)\s+(?=\[\^)", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\s+([,.;:!?])", r"\1", t)
    t = re.sub(r"\s{2,}", " ", t).strip()
    # Keep the Introduction/Summary concise: 2–3 sentences maximum.
    t = _limit_to_n_sentences(t, n=3)

    return t



def _limit_to_n_sentences(text: str, n: int = 3) -> str:
    """Deterministically keep the first n sentences (simple punctuation-based segmentation)."""
    if not text:
        return text
    s = re.sub(r"\s+", " ", str(text)).strip()
    if not s:
        return s
    out = []
    start = 0
    i = 0
    while i < len(s) and len(out) < n:
        ch = s[i]
        if ch in '.!?':
            j = i + 1
            # absorb closing quotes/brackets
            while j < len(s) and s[j] in "'\")]}":
                j += 1
            seg = s[start:j].strip()
            if seg:
                out.append(seg)
            # skip whitespace after sentence end
            while j < len(s) and s[j].isspace():
                j += 1
            start = j
            i = j
            continue
        i += 1
    # If we never hit punctuation, treat as one sentence.
    if not out:
        out = [s]
    return " ".join(out).strip()

def enforce_intro_summary_rules(text: str) -> str:
    """Hard rules for Introduction and Summary."""
    if not text:
        return text

    t = str(text)

    t = re.sub(r"\(\s*Entry\s+\d+\s*\)", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\bEntry\s+\d+\b", "", t, flags=re.IGNORECASE)

    scaffold_patterns = [
        r"\bas\s+evidenced\s+by\b",
        r"\bas\s+evident\s+in\b",
        r"\bas\s+highlighted\s+in\b",
        r"\bas\s+exemplified\s+by\b",
        r"\bas\s+illustrated\s+by\b",
        r"\billustrated\s+by\b",
        r"\bdemonstrated\s+by\b",
        r"\bas\s+seen\s+in\b",
        r"\bas\s+noted\s+in\b",
        r"\bmarked\s+by\b",
        r"\bactivities?\s+marked\s+by\b",
        r"\bactivities?\b",
        r"\binitiatives?\b",
        r"\befforts?\b",
        r"\bunique\s+identifiers?\b",
        r"\bthe\s+above\s+activities?\b",
    ]

    t = re.sub(
        r"\b(For\s+instance|For\s+example)\b\s*,?\s*(?=(?:.{0,60}\[\^))",
        "",
        t,
        flags=re.IGNORECASE
    )

    t = re.sub(r"\bsuch\s+as\b\s*(?=(?:.{0,60}\[\^))", "", t, flags=re.IGNORECASE)

    for p in scaffold_patterns:
        t = re.sub(p, "", t, flags=re.IGNORECASE)

    t = compact_consecutive_citations(t)
    t = normalize_uid_marker_placement(t)
    t = ensure_period_after_citations(t)

    # Intro/Summary must be citation-free: strip any remaining UID footnote markers.
    t = re.sub(r"\[\^\s*[A-Za-z0-9._-]+\s*\]", "", t)

    t = re.sub(
        r"(\[\^\s*[A-Za-z0-9._-]+\s*\])\s+(?=[a-z])",
        r"\1 ",
        t
    )

    t = re.sub(r"\s+", " ", t).strip()
    t = re.sub(r"\s+([,.;:!?])", r"\1", t)
    t = re.sub(r",\s*,+", ", ", t)
    t = re.sub(r"\(\s*\)", "", t)
    t = re.sub(r"\s{2,}", " ", t).strip()

    return t


def fallback_paragraph_from_card(card: dict) -> str:
    """Deterministic fallback narrative when the LLM output is missing or unusable."""
    if not card:
        return ""
    title = (card.get("Activity Name") or card.get("Activity Title") or "").strip()
    desc = (card.get("Activity Description") or "").strip()
    lead = (card.get("Lead ICO") or card.get("Submitting ICO") or "").strip()
    collab = (card.get("Collaborating ICOs/Agencies/Orgs") or "").strip()

    sentences = []
    if title and title != "—":
        sentences.append(f"{title}.")
    if desc and desc != "—":
        d = re.sub(r"\s+", " ", desc).strip()
        sentences.append(d)
    if lead and lead != "—":
        if collab and collab != "—":
            sentences.append(f"This activity is led by {lead} with collaboration from {collab}.")
        else:
            sentences.append(f"This activity is led by {lead}.")

    out = " ".join(sentences).strip()
    out = hard_clean_generated_text(out)
    out = re.sub(r"\s+", " ", out).strip()
    out = postprocess_narrative(out)
    return out


def resolve_map(cols):
    lower = {str(c).strip().lower(): str(c).strip() for c in cols}
    m = {}
    for want in CANON:
        w = want.lower()
        if w in lower:
            m[want] = lower[w]
            continue
        found = None
        for k in lower:
            if w.replace(" ", "") in k.replace(" ", ""):
                found = lower[k]
                break
        if found:
            m[want] = found
    return m


def as_str(x):
    if pd.isna(x):
        return "—"
    s = str(x)
    return s if s.strip() else "—"


def split_urls(s):
    if not s or s == "—":
        return []
    out = re.split(r"[;\s,]+", str(s).strip())
    return [p for p in out if p.lower().startswith("http")]


def build_activity_brief(uid: str, card: dict) -> str:
    """Compact, grounded brief for one activity, used to ground LLM outputs."""
    title = (card.get("Title") or card.get("Activity Title") or card.get("Headline") or "").strip()
    lead = (card.get("Lead ICO") or card.get("Lead IC") or card.get("Lead Institute") or "").strip()
    ico = (card.get("ICO") or card.get("ICO(s)") or card.get("Institute/Center") or "").strip()
    pmids = split_pmids(card.get("PMID(s)", ""))
    urls = split_urls(card.get("Web address(es)", ""))
    ref = ""
    if pmids:
        ref = f"PMID {pmids[0]}"
    elif urls:
        ref = urls[0]

    bits = [uid]
    if title:
        bits.append(f"Title: {title}")
    if lead:
        bits.append(f"Lead: {lead}")
    elif ico:
        bits.append(f"ICO: {ico}")
    if ref:
        bits.append(f"Ref: {ref}")
    return " | ".join(bits)


def build_portfolio_evidence_brief(df: pd.DataFrame, uid_index: dict, section_to_uids: dict, max_examples_per_section: int = 3) -> str:
    """Grounded evidence brief for data-driven Intro/Summary/Section overviews (offline-only)."""
    total = len(df) if df is not None else 0

    lines = [f"Portfolio size (filtered activities): {total}."]

    fy_cols = [c for c in df.columns if str(c).strip().lower() in {"fy", "fiscal year", "fiscal_year"}]
    if fy_cols:
        vals = sorted({str(v).strip() for v in df[fy_cols[0]].dropna().tolist() if str(v).strip()})
        if vals:
            lines.append("Fiscal years: " + ", ".join(vals[:6]) + ".")

    ic_cols = [c for c in df.columns if str(c).strip().lower() in {"ico", "lead ico", "lead ic", "institute", "institute/center"}]
    if ic_cols:
        s = df[ic_cols[0]].fillna("").astype(str).str.strip()
        top = s[s != ""].value_counts().head(8)
        if len(top) > 0:
            lines.append("Top contributing ICs: " + "; ".join([f"{k} ({v})" for k, v in top.items()]) + ".")

    lines.append("Representative activities by section (UID | Title | Lead/ICO | PMID/URL):")
    for sec in SECTION_ORDER:
        uids = section_to_uids.get(sec, []) or []
        if not uids:
            continue
        lines.append(f"{sec}:")
        for uid in uids[:max_examples_per_section]:
            lines.append("  - " + build_activity_brief(uid, uid_index.get(uid, {}) or {}))

    return "\n\n".join(lines).strip()


def make_authoritative_style_constraints() -> str:
    return (
        "Constraints:\n"
        "- Write in an authoritative NIH editorial style. Prefer concrete facts over generic claims.\n"
        "- Never include empty brackets like []. Do not use placeholders.\n"
        "- Do not mention 'UID' in prose.\n"
        "- When giving examples, refer to a specific activity by its Title (do not use UID in prose).\n"
        "- Avoid repetitive sentence openings; vary sentence structure.\n"
        "- No inline URLs in prose.\n"
        "- Do not use NIH entities as nouns with a leading article. Write 'NIH', not 'the NIH'; 'NCI', not 'the NCI'.\n"
        "- Do not write out NIH Institute and Center (ICO) names in full in narrative paragraphs; use their acronyms (e.g., NCI).\n"
    )


def build_intro_prompt(evidence_brief: str) -> list:
    evidence_brief = (evidence_brief or "").strip()
    return [
        {"role": "system", "content": "You write NIH triennial report narrative text."},
        {"role": "user",
         "content":
            "Draft the INTRODUCTION for this chapter.\n"
            + make_authoritative_style_constraints()
            + "\nUse ONLY the evidence brief below. If not in the brief, omit it.\n"
            "Write 2–3 sentences total. Define scope and why it matters; you may mention 1–2 concrete examples by title if they fit.\n"
            "Do not invent statistics. Do not include any citations, UID markers, footnotes, or bracket references.\n"
            "\nEVIDENCE BRIEF:\n" + evidence_brief
        },
    ]


def build_summary_prompt(evidence_brief: str) -> list:
    evidence_brief = (evidence_brief or "").strip()
    return [
        {"role": "system", "content": "You write NIH triennial report narrative text."},
        {"role": "user",
         "content":
            "Draft the SUMMARY for this chapter.\n"
            + make_authoritative_style_constraints()
            + "\nUse ONLY the evidence brief below.\n"
            "Write 2–3 sentences total. Capture cross-cutting themes and concrete highlights by title.\n"
            "Do not invent statistics. Do not include any citations, UID markers, footnotes, or bracket references.\n"
            "\nEVIDENCE BRIEF:\n" + evidence_brief
        },
    ]


def build_section_synthesis_prompt(section_name: str, section_uids: list[str], uid_index: dict) -> list:
    section_uids = section_uids or []
    briefs = []
    for uid in (section_uids or [])[:6]:
        briefs.append("- " + build_activity_brief(uid, uid_index.get(uid, {}) or {}))
    return [
        {"role": "system", "content": "You write NIH triennial report narrative text."},
        {"role": "user",
         "content":
            f"Draft the brief overview for the section: {section_name}.\n"
            + make_authoritative_style_constraints()
            + "\nUse ONLY the activity briefs below.\n"
            "Write 1 paragraph (~110–170 words). Focus on what unifies the section and the advances represented.\n"
            "Do not mention UIDs. Do not repeat details that belong in individual activity paragraphs.\n"
            "\nACTIVITY BRIEFS:\n" + "\n".join(briefs)
        },
    ]


def build_row_paragraph_prompt(uid: str, card: dict) -> list:
    brief = build_activity_brief(uid, card)
    return [
        {"role": "system", "content": "You write NIH triennial report narrative text."},
        {"role": "user",
         "content":
            "Write one activity paragraph for an NIH triennial report.\n"
            + make_authoritative_style_constraints()
            + "\nRequirements:\n"
            "- 110–160 words. One paragraph.\n"
            "- Start with the scientific contribution (avoid starting with 'Researchers' or 'Scientists').\n"
            "- Name the lead institute only once if needed.\n"
            "- Do not include inline URLs or empty brackets.\n"
            "- End with a complete sentence.\n"
            "- IMPORTANT (citation control): Append the token [[CITE]] at the end of ONLY those sentences that are directly supported by the activity brief. Do NOT add [[CITE]] to bridging/general context sentences.\n"
            + "\nACTIVITY BRIEF:\n" + brief
        },
    ]


def generate_row_paragraph(effective_system_text: str, card: dict, uid: str = None) -> str:
    """
    Generate one narrative paragraph for a single activity row/card.

    This is intentionally a thin wrapper around:
      - build_row_paragraph_prompt(...)
      - call_fmapi(...)
      - extract_text(...)

    It exists because the UI loop calls generate_row_paragraph(...) directly.
    """
    # Resolve UID if not explicitly provided
    if uid is None:
        uid = (card.get("UID") or card.get("uid") or card.get("Uid") or "").strip()
    if not uid:
        # Fall back to a deterministic placeholder to avoid hard crashes
        uid = "UNKNOWN_UID"

    messages = build_row_paragraph_prompt(uid, card)

    # If we have an effective system prompt, override the system message
    eff = (effective_system_text or "").strip()
    if eff:
        if messages and isinstance(messages[0], dict) and messages[0].get("role") == "system":
            messages[0]["content"] = eff
        else:
            messages = [{"role": "system", "content": eff}] + messages

    resp = call_fmapi(
        ENDPOINT,
        messages=messages,
        max_tokens=ROW_MAX_TOKENS,
        temperature=ROW_TEMP,
        retries=2,
    )
    para = extract_text(resp).strip()

    # Ensure it's a single paragraph (no blank lines)
    para = re.sub(r"\n\s*\n+", " ", para).strip()

    # Safety: ensure terminal punctuation
    if para and para[-1] not in ".!?":
        para += "."

    return para

def split_pmids(s):
    """Return candidate PMIDs (7–9 digits) from a string."""
    if not s or s == "—":
        return []
    return re.findall(r"\b(\d{7,9})\b", str(s))


def _count_words(txt: str) -> int:
    return len([w for w in re.findall(r"\b\w+\b", txt)])


def _split_paragraphs(md: str):
    return [p.strip() for p in re.split(r"\n\s*\n", md) if p.strip()]


def _intro_meets_shape(md: str, min_paras: int, min_words: int, target_max: int) -> bool:
    paras = _split_paragraphs(md)
    if len(paras) < min_paras:
        return False
    if len(paras) > target_max:
        return False
    for p in paras:
        if _count_words(p) < min_words:
            return False
    return True

def _summary_meets_shape(md: str, min_paras: int, min_words: int, target_max: int) -> bool:
    paras = _split_paragraphs(md)
    if len(paras) < min_paras:
        return False
    if len(paras) > target_max:
        return False
    for p in paras:
        if _count_words(p) < min_words:
            return False
    return True


def build_footnotes_from_uid_markers(md_text: str, uid_index: dict) -> tuple[str, str, str]:
    """
    Convert UID markers [^UID] into stable numeric footnotes (deduplicated per UID),
    and build a References section that lists EVERY footnote source in the SAME order.
    """
    UID_MARK_RE = re.compile(r"\[\^\s*([A-Za-z0-9._-]+)\s*\]")
    if not md_text:
        return "", "", ""

    known_uids = set(uid_index.keys())

    def _drop_unknown(m: re.Match) -> str:
        uid = m.group(1)
        return m.group(0) if uid in known_uids else ""

    md_clean = UID_MARK_RE.sub(_drop_unknown, md_text)

    def _split_urls(cell: str) -> list[str]:
        if cell is None:
            return []
        s = str(cell).strip()
        if not s or s.lower() in ("nan", "none", "—"):
            return []
        parts = re.split(r"[;\n]+", s)
        urls: list[str] = []
        for part in parts:
            part = part.strip()
            if not part:
                continue
            found = re.findall(r"https?://\S+", part)
            for u in found:
                u = u.rstrip(").,;]")
                urls.append(u)
        return urls

    def _extract_pmids(cell: str) -> list[str]:
        if cell is None:
            return []
        s = str(cell)
        pmids = re.findall(r"\b\d{6,9}\b", s)
        out: list[str] = []
        seen = set()
        for p in pmids:
            if p not in seen:
                seen.add(p)
                out.append(p)
        return out

    def _canonicalize_url(url: str) -> str:
        u = (url or "").strip()
        u = u.rstrip(").,;]")
        if len(u) > 8 and u.endswith("/"):
            u = u[:-1]
        return u

    def _make_ref_text(uid: str) -> str:
        row = uid_index.get(uid, {}) or {}
        web_urls = _split_urls(row.get("Web address(es)") or "")
        pmids = _extract_pmids(row.get("PMID(s)") or "")

        # --- PMID → NLM (PubMed) reference formatting ---
        # If PMID metadata cannot be fetched/parsed, fall back to a stable PubMed URL.
        @lru_cache(maxsize=2048)
        def _pmid_to_nlm(pmid: str) -> Optional[str]:
            pmid = (pmid or "").strip()
            if not pmid:
                return None

            # NCBI E-utilities (no API key required for low volume; add key if needed later)
            url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
            params = {"db": "pubmed", "id": pmid, "retmode": "xml"}
            try:
                r = requests.get(url, params=params, timeout=20)
                r.raise_for_status()
                xml_text = r.text
            except Exception:
                return None

            try:
                root = ET.fromstring(xml_text)
            except Exception:
                return None

            def _t(x: Optional[str]) -> str:
                return (x or "").strip()

            # First PubmedArticle only
            art = root.find(".//PubmedArticle")
            if art is None:
                return None

            # Authors
            authors = []
            for a in art.findall(".//Article/AuthorList/Author"):
                last = _t(a.findtext("LastName"))
                initials = _t(a.findtext("Initials"))
                coll = _t(a.findtext("CollectiveName"))
                if coll:
                    authors.append(coll)
                elif last and initials:
                    authors.append(f"{last} {initials}")
                elif last:
                    authors.append(last)

            if authors:
                if len(authors) > 6:
                    authors_txt = ", ".join(authors[:6]) + ", et al"
                else:
                    authors_txt = ", ".join(authors)
                authors_txt = authors_txt.rstrip(".") + "."
            else:
                authors_txt = ""

            # Title
            title = _t(art.findtext(".//Article/ArticleTitle"))
            title = re.sub(r"\s+", " ", title).strip()
            if title and not title.endswith("."):
                title += "."
            # Journal
            journal = _t(art.findtext(".//Article/Journal/ISOAbbreviation")) or _t(art.findtext(".//Article/Journal/Title"))
            if journal and not journal.endswith("."):
                journal += "."

            # Pub date
            year = _t(art.findtext(".//Article/Journal/JournalIssue/PubDate/Year")) or _t(art.findtext(".//ArticleDate/Year"))
            month = _t(art.findtext(".//Article/Journal/JournalIssue/PubDate/Month")) or _t(art.findtext(".//ArticleDate/Month"))
            day = _t(art.findtext(".//Article/Journal/JournalIssue/PubDate/Day")) or _t(art.findtext(".//ArticleDate/Day"))

            def _month_norm(m: str) -> str:
                m = (m or "").strip()
                if not m:
                    return ""
                # PubMed sometimes returns numeric month or 3-letter abbreviation
                if m.isdigit():
                    mi = int(m)
                    if 1 <= mi <= 12:
                        return calendar.month_abbr[mi]
                    return ""
                m3 = m[:3].title()
                # Accept already-abbreviated forms (Jan, Feb, Mar, ...)
                if m3 in list(calendar.month_abbr):
                    return m3
                return m3

            month = _month_norm(month)
            date_bits = [b for b in [year, month, day] if b]
            date_txt = (" ".join(date_bits) + ";") if date_bits else ""

            # Volume/Issue/Pages
            vol = _t(art.findtext(".//Article/Journal/JournalIssue/Volume"))
            iss = _t(art.findtext(".//Article/Journal/JournalIssue/Issue"))
            pages = _t(art.findtext(".//Article/Pagination/MedlinePgn"))
            vip = ""
            if vol:
                vip += vol
                if iss:
                    vip += f"({iss})"
            if pages:
                vip += f":{pages}"
            if vip and not vip.endswith("."):
                vip += "."

            # DOI
            doi = ""
            for aid in art.findall(".//ArticleIdList/ArticleId"):
                if (aid.get("IdType") or "").lower() == "doi":
                    doi = _t(aid.text)
                    break
            doi_txt = f" doi: {doi}." if doi else ""

            # PMID
            pmid_txt = f" PMID: {pmid}."

            # NLM-ish assembly (concise, journal-article style)
            parts = [p for p in [authors_txt, title, journal, date_txt + (vip or "")] if p]
            if not parts:
                return None
            core = " ".join([p.strip() for p in parts]).strip()
            # Ensure single spaces and clean punctuation spacing
            core = re.sub(r"\s+", " ", core)
            core = re.sub(r"\s+([,.;:])", r"\1", core).strip()

            return (core + doi_txt + pmid_txt).strip()

        if pmids:
            pmid = pmids[0]
            nlm = _pmid_to_nlm(pmid)
            if nlm:
                return nlm

            # Fallback if PubMed metadata is unreachable
            pubmed_url = _canonicalize_url(f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/")
            return f"PMID: {pmid}. {pubmed_url}"

        if web_urls:
            url = _canonicalize_url(web_urls[0])
            if url:
                return url

        return "Source unavailable"


    footnotes: list[tuple[str, str]] = []
    references_lines: list[str] = []

    out_parts: list[str] = []
    last = 0
    fn_counter = 0
    uid_to_fn: dict[str, tuple[int, str]] = {}  # uid -> (num, fn_label)

    for m in UID_MARK_RE.finditer(md_clean):
        uid = m.group(1)

        out_parts.append(md_clean[last:m.start()])
        last = m.end()

        if uid not in known_uids:
            continue

        if uid not in uid_to_fn:
            fn_counter += 1
            fn_label = f"fn{fn_counter}"
            ref_text = _make_ref_text(uid).strip()
            uid_to_fn[uid] = (fn_counter, fn_label)

            footnotes.append((fn_label, ref_text))
            references_lines.append(f"{fn_counter} {ref_text}")

        _, fn_label = uid_to_fn[uid]
        out_parts.append(f"[^{fn_label}]")

    out_parts.append(md_clean[last:])
    md_with_numeric = "".join(out_parts)

    footnote_lines: list[str] = []
    for fn_label, fn_text in footnotes:
        footnote_lines.append(f"[^{fn_label}]: {fn_text}")
    footnote_block = "\n".join(footnote_lines).strip() + ("\n" if footnote_lines else "")

    references_md = "\n\n".join(references_lines).strip() + ("\n" if references_lines else "")
    
    return md_with_numeric, footnote_block, references_md


def _inject_references_section(md_text: str, references_md: str) -> str:
    """Insert references list under the '## References' heading if present."""
    if not md_text or not references_md:
        return md_text

    refs_re = re.compile(r"##\s*References\s*\n", re.I)
    if not refs_re.search(md_text):
        return md_text

    return refs_re.sub(f"## References\n\n{references_md}\n\n", md_text)


# -----------------------------
# Acronyms extraction + section (DETERMINISTIC)
# -----------------------------
_ACRONYM_TOKEN_RE = re.compile(r"\b[A-Z][A-Z0-9]{1,9}(?:-[A-Z0-9]{1,10})?\b")

_ACRONYM_STOP = {
    "UID", "PMID", "DOI", "URL", "URLs", "U.S", "US", "USA", "FY", "FYS",
    "ICO", "ICOs", "IC", "ICs",
}

_ICO_ACRONYMS = {
    "NIH", "NIAMS", "NCI", "NCATS", "NIAID", "NHLBI", "NIDDK", "NIA", "NIMH", "NINDS", "NICHD",
    "NIEHS", "NIAAA", "NIBIB", "NIGMS", "NHGRI", "NLM", "OD",
}

_ACRONYM_EXPANSIONS = {
    "AI": "Artificial intelligence",
    "ML": "Machine learning",
    "IL-5": "Interleukin-5",
    "NTP": "National Toxicology Program",
    "TK6": "Human TK6 lymphoblastoid cell line",
    "NCI": "National Cancer Institute",
    "NIH": "National Institutes of Health",
    "NIEHS": "National Institute of Environmental Health Sciences",
    "NCATS": "National Center for Advancing Translational Sciences",
    "NINDS": "National Institute of Neurological Disorders and Stroke",
    "NICHD": "Eunice Kennedy Shriver National Institute of Child Health and Human Development",
    "NIAMS": "National Institute of Arthritis and Musculoskeletal and Skin Diseases",
    "NHLBI": "National Heart, Lung, and Blood Institute",
    "NIAID": "National Institute of Allergy and Infectious Diseases",
    "OSC": "Office of Strategic Coordination",
}

_ICO_EXPANSIONS = {
    "NIH": "National Institutes of Health",
    "NIAMS": "National Institute of Arthritis and Musculoskeletal and Skin Diseases",
    "NCI": "National Cancer Institute",
    "NCATS": "National Center for Advancing Translational Sciences",
    "NIAID": "National Institute of Allergy and Infectious Diseases",
    "NHLBI": "National Heart, Lung, and Blood Institute",
    "NIDDK": "National Institute of Diabetes and Digestive and Kidney Diseases",
    "NIA": "National Institute on Aging",
    "NIMH": "National Institute of Mental Health",
    "NINDS": "National Institute of Neurological Disorders and Stroke",
    "NICHD": "Eunice Kennedy Shriver National Institute of Child Health and Human Development",
    "NIEHS": "National Institute of Environmental Health Sciences",
    "NIAAA": "National Institute on Alcohol Abuse and Alcoholism",
    "NIBIB": "National Institute of Biomedical Imaging and Bioengineering",
    "NIGMS": "National Institute of General Medical Sciences",
    "NHGRI": "National Human Genome Research Institute",
    "NLM": "National Library of Medicine",
    "OD": "Office of the Director",
}

_ACRONYM_EXPANSIONS.update(_ICO_EXPANSIONS)

def _normalize_acronym_token(tok: str) -> str:
    tok = (tok or "").strip()
    tok = tok.rstrip(".,;:)]}")
    tok = tok.lstrip("([{")
    return tok

def _extract_acronyms_from_blob(text: str) -> list[str]:
    if not text:
        return []
    cleaned = re.sub(r"https?://\S+", " ", text)
    found = []
    for m in _ACRONYM_TOKEN_RE.finditer(cleaned):
        tok = _normalize_acronym_token(m.group(0))
        if not tok:
            continue
        if tok in _ACRONYM_STOP:
            continue
        if re.fullmatch(r"\d+", tok):
            continue
        found.append(tok)
    return found

def _extract_explicit_acronyms_from_cards(cards: list | None) -> dict:
    if not cards:
        return {}

    possible_keys = ["Acronyms", "Acronym", "Acronym(s)", "Abbreviation(s)", "Abbreviations"]
    out = {}

    for c in cards:
        for k in possible_keys:
            raw = c.get(k)
            if raw is None:
                continue
            s = str(raw).strip()
            if not s or s in ("—", "nan", "None"):
                continue

            chunks = re.split(r"[;\n]+", s)
            for ch in chunks:
                ch = ch.strip()
                if not ch:
                    continue

                if "=" in ch:
                    a, b = [x.strip() for x in ch.split("=", 1)]
                elif " — " in ch:
                    a, b = [x.strip() for x in ch.split(" — ", 1)]
                elif " - " in ch:
                    a, b = [x.strip() for x in ch.split(" - ", 1)]
                else:
                    a, b = ch.strip(), ""

                a = _normalize_acronym_token(a)
                if not a:
                    continue
                if not re.fullmatch(r"[A-Z0-9]{2,10}(?:-[A-Z0-9]{2,10})?", a):
                    continue

                out[a] = b.strip()

    return out


def _extract_pmids_from_cards(cards: list | None, max_pmids: int = 80) -> list[str]:
    """Collect PMIDs from cards for PubMed-based acronym expansion inference."""
    if not cards:
        return []
    out: list[str] = []
    seen = set()
    for c in cards:
        for k in ("PMID(s)", "PMID", "PMIDs", "PubMed ID", "PubMed IDs"):
            raw = c.get(k)
            if raw is None:
                continue
            for p in re.findall(r"\b(\d{7,9})\b", str(raw)):
                if p not in seen:
                    seen.add(p)
                    out.append(p)
                    if len(out) >= max_pmids:
                        return out
    return out


@lru_cache(maxsize=2048)
def _pmid_to_pubmed_text(pmid: str) -> str:
    """
    Fetch Title + Abstract text for a PMID via NCBI E-utilities (PubMed XML).
    Returns an empty string if anything fails.
    """
    pmid = (pmid or "").strip()
    if not pmid:
        return ""
    url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
    params = {"db": "pubmed", "id": pmid, "retmode": "xml"}
    try:
        r = requests.get(url, params=params, timeout=20)
        r.raise_for_status()
    except Exception:
        return ""

    try:
        root = ET.fromstring(r.text)
    except Exception:
        return ""

    art = root.find(".//PubmedArticle")
    if art is None:
        return ""

    title_node = art.find(".//ArticleTitle")
    title = "".join(title_node.itertext()).strip() if title_node is not None else ""
    # Abstract may have multiple AbstractText nodes
    abst_parts = []
    for a in art.findall(".//Abstract/AbstractText"):
        abst_parts.append("".join(a.itertext()))
    abstract = " ".join([p.strip() for p in abst_parts if p and p.strip()])

    blob = " ".join([str(title).strip(), str(abstract).strip()]).strip()
    return re.sub(r"\s+", " ", blob)


def expand_acronyms_from_pubmed(acronyms: list[str], pmids: list[str]) -> dict:
    """
    Infer acronym expansions by mining Title/Abstract text from PubMed articles
    (NLM/NCBI source) for patterns like:
      - 'long form (ACR)'
      - 'ACR (long form)'

    We only use the PMIDs already present in the dataset/cards (to keep it relevant).
    """
    acronyms = [a for a in (acronyms or []) if a and re.fullmatch(r"[A-Z0-9]{2,10}(?:-[A-Z0-9]{2,10})?", a)]
    acronyms = sorted(set(acronyms))
    if not acronyms or not pmids:
        return {}

    # Limit network work
    pmids = [p for p in pmids if p and re.fullmatch(r"\d{7,9}", str(p).strip())]
    pmids = pmids[:50]

    # Pre-fetch blobs
    blobs = []
    for p in pmids:
        t = _pmid_to_pubmed_text(str(p))
        if t:
            blobs.append(t)

    if not blobs:
        return {}

    # Build candidate expansions per acronym
    best: dict[str, str] = {}
    for acr in acronyms:
        # Capture 4–80 chars of words/spaces/hyphens before/after
        # Example: "Childhood Cancer Data Initiative (CCDI)"
        # Example: "CCDI (Childhood Cancer Data Initiative)"
        re_after = re.compile(rf"\b{re.escape(acr)}\s*\(\s*([A-Za-z][A-Za-z0-9\-/,&' ]{{4,80}}?)\s*\)")
        re_before = re.compile(rf"\b([A-Za-z][A-Za-z0-9\-/,&' ]{{4,80}}?)\s*\(\s*{re.escape(acr)}\s*\)")

        counts: dict[str, int] = {}
        for blob in blobs:
            for m in re_before.finditer(blob):
                cand = re.sub(r"\s+", " ", m.group(1)).strip(" .;:,")
                if _validate_acronym_expansion(acr, cand):
                    counts[cand] = counts.get(cand, 0) + 1
            for m in re_after.finditer(blob):
                cand = re.sub(r"\s+", " ", m.group(1)).strip(" .;:,")
                if _validate_acronym_expansion(acr, cand):
                    counts[cand] = counts.get(cand, 0) + 1

        if counts:
            # Deterministic tie-break: highest frequency, then shortest, then alpha
            best_cand = sorted(counts.items(), key=lambda kv: (-kv[1], len(kv[0]), kv[0].lower()))[0][0]
            best[acr] = best_cand

    return best

def build_acronyms_section(md_text: str, cards: list | None = None, system_text: str | None = None) -> str:
    md_text = md_text or ""

    explicit = _extract_explicit_acronyms_from_cards(cards)
    inferred = set(_extract_acronyms_from_blob(md_text))

    all_acrs = set(explicit.keys()) | inferred

    ico_exclude = set(_ICO_ACRONYMS)
    try:
        ico_exclude |= set(_ICO_FULLNAME_TO_ACR.values())
    except Exception:
        pass
    all_acrs = {a for a in all_acrs if a not in ico_exclude}

    if not all_acrs:
        return ""

    ordered = sorted(all_acrs)

    missing: list[str] = []
    for acr in ordered:
        exp = (explicit.get(acr) or "").strip()
        if not exp:
            exp = (_ACRONYM_EXPANSIONS.get(acr, "") or "").strip()
        if not exp:
            missing.append(acr)
    pubmed_map = {}
    if missing:
        try:
            pmids = _extract_pmids_from_cards(cards)
            pubmed_map = expand_acronyms_from_pubmed(missing, pmids)
        except Exception:
            pubmed_map = {}

    llm_map = {}
    remaining = [a for a in missing if not (pubmed_map.get(a) or "").strip()]
    if remaining:
        try:
            llm_map = expand_acronyms_with_llm(
                system_text=system_text_base if "system_text_base" in globals() else "You write NIH triennial report narrative text.",
                acronyms=remaining,
                context_text=md_text,
                max_tokens=450,
            )
        except Exception:
            llm_map = {}

    lines: list[str] = []
    for acr in ordered:
        exp = (explicit.get(acr) or "").strip()
        if not exp:
            exp = (_ACRONYM_EXPANSIONS.get(acr, "") or "").strip()
        if not exp:
            exp = (pubmed_map.get(acr) or "").strip()
        if not exp:
            exp = (llm_map.get(acr) or "").strip()

        # Remove low-quality placeholders
        if exp:
            exp_norm = exp.strip().strip(".")
            if exp_norm.lower() == acr.strip().lower():
                exp = ""
            elif exp_norm.lower() in {"expansion not specified", "not specified", "unknown", "n/a"}:
                exp = ""

        # If we still can't infer an expansion, omit the acronym rather than printing a placeholder.
        if not exp:
            continue

        lines.append(f"**{acr}** — {exp.strip().strip('.')}")
    return "\n\n".join(lines).strip()



def _validate_acronym_expansion(acr: str, exp: str) -> bool:
    if not exp:
        return False
    exp = exp.strip()
    if exp.lower() in {"unknown", "n/a", "not sure", "unsure", "not specified"}:
        return False
    if len(exp) < 6:
        return False
    if acr.lower() in exp.lower():
        return False
    if re.fullmatch(r"[\W_]+", exp):
        return False
    return True


def expand_acronyms_with_llm(system_text: str, acronyms: list[str], context_text: str, max_tokens: int = 450) -> dict:
    acronyms = [a for a in acronyms if a and re.fullmatch(r"[A-Z0-9]{2,10}(?:-[A-Z0-9]{2,10})?", a)]
    acronyms = sorted(set(acronyms))
    if not acronyms:
        return {}

    ctx = (context_text or "").strip()
    if len(ctx) > 6000:
        ctx = ctx[:6000]

    instr = (
        "You are expanding acronyms for an NIH-style report.\n"
        "Rules:\n"
        "1) Expand ONLY the acronyms in the provided list. Do not add new keys.\n"
        "2) If you are not highly confident, use null for expansion.\n"
        "3) Return JSON only, no markdown fences.\n"
        'Output schema: {"ACR": {"expansion": "Full term", "confidence": 0.0}}\n'
        "4) Use the context text to disambiguate.\n"
    )

    payload = {
        "acronyms": acronyms,
        "context": ctx,
    }

    resp = call_fmapi(
        ENDPOINT,
        messages=[
            {"role": "system", "content": system_text},
            {"role": "user", "content": instr + "\n" + json.dumps(payload, ensure_ascii=False)},
        ],
        max_tokens=max_tokens,
        temperature=0.0,
    )

    raw = extract_text(resp).strip()

    try:
        data = json.loads(raw)
    except Exception:
        return {}

    out = {}
    for acr in acronyms:
        v = data.get(acr)
        if not isinstance(v, dict):
            continue
        exp = v.get("expansion")
        conf = v.get("confidence", 0.0)
        if exp is None:
            continue
        exp = str(exp).strip()
        try:
            conf_f = float(conf)
        except Exception:
            conf_f = 0.0

        if conf_f < 0.70:
            continue

        if _validate_acronym_expansion(acr, exp):
            out[acr] = exp

    return out

def apply_primary_color_to_docx(docx_path: str, rgb: Tuple[int, int, int]) -> None:
    if not DOCX_AVAILABLE:
        return

    r, g, b = rgb
    color = RGBColor(r, g, b)

    doc = Document(docx_path)

    for style in doc.styles:
        try:
            if hasattr(style, "font") and style.font is not None:
                style.font.color.rgb = color
        except Exception:
            pass

    def _apply_runs_in_paragraph(par):
        for run in par.runs:
            try:
                run.font.color.rgb = color
            except Exception:
                pass

    def _apply_in_table(tbl):
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_runs_in_paragraph(p)
                for nested in cell.tables:
                    _apply_in_table(nested)

    for p in doc.paragraphs:
        _apply_runs_in_paragraph(p)

    for tbl in doc.tables:
        _apply_in_table(tbl)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _apply_runs_in_paragraph(p)
        for p in section.footer.paragraphs:
            _apply_runs_in_paragraph(p)

        for tbl in section.header.tables:
            _apply_in_table(tbl)
        for tbl in section.footer.tables:
            _apply_in_table(tbl)

    doc.save(docx_path)


# =============================
# 5) Load style prompt + Excel (cached)
# =============================
@st.cache_data(show_spinner=False)
def load_system_text(style_prompt_path_local: str) -> str:
    p = Path(style_prompt_path_local)
    if not p.exists():
        raise FileNotFoundError(f"Missing style prompt (local staged): {style_prompt_path_local}")
    sysj = json.loads(p.read_text(encoding="utf-8"))
    content = sysj.get("content", "")
    return "\n".join(content) if isinstance(content, list) else str(content)


def build_effective_system_text(system_text_base: str, style_override: str) -> str:
    """
    Combine the base system prompt from style_prompt.json with an optional per-run override.

    The override must NEVER leak internal implementation details into the user-visible plan.
    This function only concatenates text; downstream code is responsible for scrub/filters.
    """
    base = (system_text_base or "").strip()
    override = (style_override or "").strip()
    if not override:
        return base
    # Keep formatting predictable for model + routing heuristics.
    return (base + "\n\n" if base else "") + "Run-specific style override:\n" + override


@st.cache_data(show_spinner=False)
def load_excel(excel_path_local: str) -> pd.DataFrame:
    p = Path(excel_path_local)
    if not p.exists():
        raise FileNotFoundError(f"Missing Excel (local staged): {excel_path_local}")
    df = pd.read_excel(excel_path_local, sheet_name=0)
    df.columns = [str(c).strip() for c in df.columns]
    return df


def resolve_column(df: pd.DataFrame, col_name: str) -> str:
    if col_name in df.columns:
        return col_name
    lc = {c.lower(): c for c in df.columns}
    return lc.get(col_name.lower(), col_name)

def resolve_any_column(df: pd.DataFrame, candidates: list[str], fallback: str) -> str:
    """Return the first candidate column that exists (case-insensitive), else fallback."""
    cols = list(df.columns)
    lc = {str(c).strip().lower(): str(c).strip() for c in cols}
    for cand in candidates:
        if cand in cols:
            return cand
        hit = lc.get(str(cand).strip().lower())
        if hit:
            return hit
    return resolve_column(df, fallback)


def dropdown_values(df: pd.DataFrame, col: str) -> list[str]:
    vals = df[col].dropna().astype(str).map(str.strip)
    vals = vals[vals != ""]
    return sorted(set(vals.tolist()))


# =============================
# 6) Model Serving call
# =============================
_FENCE = re.compile(r"```(?:json|md)?\s*|```", re.I)
_URL = re.compile(r"https?://\S+")
_EMPTY_BRACKETS = re.compile(r"\[\s*\]")
_ODD_SUP = re.compile(r"[\u2070-\u209F\u02B0-\u02FF]")

def extract_text(d) -> str:
    msg = None
    try:
        msg = d["choices"][0]["message"]["content"]
    except Exception:
        pass

    if isinstance(msg, list):
        parts = []
        for chunk in msg:
            if isinstance(chunk, dict) and chunk.get("type") == "reasoning":
                continue
            if isinstance(chunk, dict) and chunk.get("type") == "text":
                parts.append(chunk.get("text", ""))
        txt = "\n".join(parts)
    elif isinstance(msg, str):
        txt = msg
    else:
        txt = json.dumps(d, indent=2)

    txt = _FENCE.sub("", txt)
    txt = _URL.sub("", txt)
    txt = _ODD_SUP.sub("", txt)
    txt = re.sub(r"[ \t]+\n", "\n", txt)
    return txt.strip()


def hard_clean_generated_text(txt: str) -> str:
    if txt is None:
        return ""
    s = str(txt)
    s = re.sub(r"```(?:json|md|markdown|text)?\s*|```", "", s, flags=re.I)
    s = re.sub(r"^\s{0,3}#{1,6}\s+.*$", "", s, flags=re.MULTILINE)
    s = re.sub(r"https?://\S+", "", s)
    s = re.sub(r"[\u2070-\u209F\u02B0-\u02FF]", "", s)
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()

def call_fmapi(endpoint: str, messages, max_tokens: int, temperature: float, retries: int = 2):
    host = get_workspace_host().rstrip("/")
    candidate_urls = [
        f"{host}/api/2.0/serving-endpoints/{endpoint}/invocations",
        f"{host}/serving-endpoints/{endpoint}/invocations",
    ]

    payload = {"messages": messages, "max_tokens": max_tokens, "temperature": temperature}
    last_err = None

    for attempt in range(retries + 1):
        for url in candidate_urls:
            try:
                headers = {**auth_headers(), "Content-Type": "application/json"}
                r = requests.post(url, headers=headers, json=payload, timeout=180)

                if r.status_code == 401 and attempt == 0:
                    try:
                        get_oauth_token.clear()
                    except Exception:
                        pass
                    headers = {**auth_headers(), "Content-Type": "application/json"}
                    r = requests.post(url, headers=headers, json=payload, timeout=180)

                if r.status_code == 404:
                    continue

                r.raise_for_status()
                return r.json()

            except Exception as e:
                last_err = e

        time.sleep(0.8)

    raise RuntimeError(
        f"Serving invocation failed after retries. Last error: {last_err}. "
        f"Tried URLs: {candidate_urls}"
    )


# =============================
# 6.1) LLM Narrator
# =============================
NARRATOR_MAX_TOKENS = 60
NARRATOR_TEMP = 0.35

NARRATOR_SYSTEM = (
    "You are a concise progress narrator inside a report-generation app.\n"
    "Write exactly ONE short sentence describing what is happening right now.\n"
    "Constraints:\n"
    "- No bullet points, no headings.\n"
    "- No URLs.\n"
    "- No quotes.\n"
    "- No emojis.\n"
    "- Keep it under 18 words.\n"
    "- Businesslike tone.\n"
)


def narrator_line(stage: str, detail: str, context: dict) -> str:
    payload = {
        "stage": stage,
        "detail": detail,
        "field": context.get("field"),
        "counts": context.get("counts", {}),
    }
    messages = [
        {"role": "system", "content": NARRATOR_SYSTEM},
        {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
    ]
    resp = call_fmapi(
        ENDPOINT,
        messages=messages,
        max_tokens=NARRATOR_MAX_TOKENS,
        temperature=NARRATOR_TEMP,
    )
    txt = extract_text(resp)
    txt = re.sub(r"\s+", " ", txt).strip()
    words = txt.split()
    if len(words) > 18:
        txt = " ".join(words[:18]).rstrip(".") + "."
    if not txt.endswith("."):
        txt += "."
    return txt


# =============================
# 6.2) Preview generator (LLM)
# =============================
PLAN_MAX_TOKENS = 650
PLAN_TEMP = 0.2

PLAN_SYSTEM = (
    "You are a planning assistant for a triennial report generator.\\n"
    "Output a short, clear preview in numbered steps.\\n"
    "Do NOT mention Pandoc, pypandoc, or any document-conversion tooling.\\n"
    "Do NOT mention internal implementation details (paths, binaries, libraries, or deployment specifics).\\n"
    "No code. No URLs. No markdown headings.\\n"
)


def _scrub_plan_preview_text(text: str) -> str:
    """Remove any Pandoc/tooling mentions from LLM plan previews (client requirement)."""
    if not text:
        return text
    lines = text.splitlines()
    cleaned = []
    for ln in lines:
        if re.search(r"\b(pandoc|pypandoc)\b", ln, flags=re.IGNORECASE):
            continue
        cleaned.append(ln)
    out = "\n".join(cleaned).strip()
    # If the model inlined the word mid-sentence, scrub it safely.
    out = re.sub(r"\b(pandoc|pypandoc)\b", "", out, flags=re.IGNORECASE)
    out = re.sub(r"\s{2,}", " ", out)
    return out


def generate_plan(field_value: str, uid_list: list[str], counts: dict, style_override: str) -> str:
    payload = {
        "field": field_value,
        "counts": counts,
        "uids_preview": uid_list[:30],
        "sections": SECTION_ORDER,
        "style_override": style_override.strip() if style_override else "",
        "pipeline": [
            "Filter rows",
            "Build cards and UID index",
            "Generate row paragraphs (each ends with UID marker)",
            "Route UIDs into sections",
            "Generate Summary",
            "Generate Introduction",
            "Generate section syntheses",
            "Assemble markdown",
            "Convert UID markers to numeric footnotes",
            "Render the final report document",
            "Optional: apply DOCX primary color override",
            "Publish to DBFS FileStore and offer download",
        ],
    }
    messages = [
        {"role": "system", "content": PLAN_SYSTEM},
        {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
    ]
    resp = call_fmapi(ENDPOINT, messages=messages, max_tokens=PLAN_MAX_TOKENS, temperature=PLAN_TEMP)
    return _scrub_plan_preview_text(extract_text(resp))


# =============================
# 7) Core pipeline functions
# =============================
def make_card(row, cmap: dict) -> dict:
    card = {k: as_str(row.get(cmap.get(k, k), "—")) for k in CANON}
    urls = split_urls(card.get("Web address(es)", ""))
    pmids = [f"https://pubmed.ncbi.nlm.nih.gov/{p}/" for p in split_pmids(card.get("PMID(s)", ""))]
    card["_citations"] = [u for u in (urls + pmids) if u]
    return card


def _split_sentences_preserve(text: str):
    """
    Pragmatic sentence splitter that preserves delimiters/spaces.
    Works with citation markers placed right after punctuation, e.g. "Sentence.[^UID] Next..."
    """
    if not text:
        return []
    sent_end = r'(?:[\.!\?](?:\[\^[^\]]+\])*(?:["\')\]\}]+)?)'
    return re.findall(rf'.*?{sent_end}(?:\s+|$)|.+$', text, flags=re.S)


def collapse_consecutive_citation_runs(paragraph: str) -> str:
    """
    Collapse consecutive/cascading citations that refer to the same reference set.

    - Footnote markers look like: [^UID]
    - If consecutive sentences carry the same marker set, keep ONE marker set at the end
      of the last sentence in that run.
    - If a sentence has NO markers, it breaks the run.
    - If marker sets change, close the previous run at the previous sentence end.
    """
    txt = (paragraph or "").strip()
    if not txt:
        return txt

    chunks = _split_sentences_preserve(txt)

    marker_re = re.compile(r'\[\^([^\]]+)\]')

    def _extract_and_strip(chunk: str):
        markers = marker_re.findall(chunk)
        seen = set()
        markers_norm = tuple([m for m in markers if not (m in seen or seen.add(m))])
        body = marker_re.sub("", chunk)
        return body, markers_norm

    def _append_markers_preserve_space(body: str, markers):
        if not markers:
            return body
        m = re.search(r'(\s*)$', body)
        trail = m.group(1) if m else ""
        core = body[:-len(trail)] if trail else body
        return core + "".join([f"[^{x}]" for x in markers]) + trail

    out = []
    prev_body = None
    prev_markers = None

    for chunk in chunks:
        body, markers_norm = _extract_and_strip(chunk)
        markers_norm = markers_norm if markers_norm else None

        if prev_body is None:
            prev_body, prev_markers = body, markers_norm
            continue

        if markers_norm and prev_markers and markers_norm == prev_markers:
            # Continue run: emit previous sentence WITHOUT markers; carry markers to the end of the run
            out.append(prev_body)
            prev_body = body
            continue

        # Close previous run
        if prev_markers:
            prev_body = _append_markers_preserve_space(prev_body, prev_markers)
        out.append(prev_body)

        # Start new run
        prev_body, prev_markers = body, markers_norm

    # Flush last sentence
    if prev_body is not None:
        if prev_markers:
            prev_body = _append_markers_preserve_space(prev_body, prev_markers)
        out.append(prev_body)

    joined = "".join(out)
    joined = re.sub(r'[ \t]{2,}', ' ', joined)
    return joined.strip()


def enforce_uid_markers_after_each_sentence(paragraph: str, uid: str) -> str:
    """
    Ensure the paragraph is citeable while also collapsing consecutive same-reference runs.

    - If the paragraph has NO markers at all, append [^<UID>] once at the end.
    - If the paragraph already has markers (including multiple references), do not
      overwrite them; just collapse consecutive runs.
    """
    txt = (paragraph or "").strip()
    if not txt:
        return txt

    if "[^" not in txt:
        if not re.search(r'[.!?]["\')\]]?\s*$', txt):
            txt = txt.rstrip() + "."
        txt = txt.rstrip() + f" [^{uid}]"

    return collapse_consecutive_citation_runs(txt)



def top_participating_ics(cards, k=8):
    counter = collections.Counter()
    for c in cards:
        for key in ("Submitting ICO", "Lead ICO"):
            val = (c.get(key) or "—").strip()
            if val and val != "—":
                counter[val] += 1
    return [name for name, _ in counter.most_common(k)]


def sanitize_intro(md: str) -> str:
    md = re.sub(r"^\s{0,3}#{1,6}\s+.*$", "", md, flags=re.MULTILINE)
    md = re.sub(r"\n{2,}", "\n\n", md)
    md = re.sub(r"[ \t]+\n", "\n", md)
    return md.strip()


def generate_intro(system_text: str, cards: list[dict], uid_index: dict, field_value: str):
    payload = {
        "meta": {"field_filter": field_value, "fiscal_years": []},
        "counts": {"rows": len(cards), "unique_uids": len(uid_index)},
        "institutes_top": top_participating_ics(cards, k=8),
        "allowed_uids": sorted(uid_index.keys()),
    }

    instr = (
        f"- Write EXACTLY 2 long, substantive paragraphs for the Introduction of an NIH Triennial report.\n"
        "- Use only the facts in the provided payload (meta, counts, top institutes, and allowed_uids).\n"
        "- Do NOT start paragraphs with 'Research', 'Research in', 'Researchers', or 'Studies'.\n"
        "- Vary paragraph openers using method-first, advance-first, or infrastructure-first openings.\n"
        "- Do NOT begin any paragraph with meta-research framing such as "
        "'By leveraging', 'Research in', 'Asthma research', or 'Studies have shown'.\n"
        "- Begin each paragraph with a concrete scientific advance, method, "
        "infrastructure, or implementation outcome.\n"
        "- Do not invent fiscal years; if fiscal_years is empty, do not mention an FY range.\n"
        "- Discuss scientific aims, collaboration patterns, infrastructure/resources, equity/access considerations, translational impact, and implementation context.\n"
        "- Do NOT introduce citations with phrases like 'as evident in', 'marked by', 'for instance', 'for example', 'such as', or 'including'.\n"
        "- Write the claim as normal prose, then place citations silently at the END of the sentence.\n"
        "- If multiple citations support one sentence, stack markers with no conjunctions, e.g., .[^UID1][^UID2]\n"
        "- Include at least four UID markers overall to anchor claims, formatted as [^<UID>] and placed sentence-terminal.\n"
        "- You may only use UID markers from allowed_uids; do not create new UIDs.\n"
        "- No bullets; output clean multi-paragraph Markdown prose.\n"
        "- No URLs/PMIDs/JSON/metadata; output pure prose paragraphs only.\n"
        f"- Each paragraph must be at least {INTRO_MIN_WORDS} words.\n"
    )

    content = instr + "\n" + json.dumps(payload, ensure_ascii=False)
    resp = call_fmapi(
        ENDPOINT,
        messages=[{"role": "system", "content": system_text}, {"role": "user", "content": content}],
        max_tokens=MAX_TOKENS_INTRO,
        temperature=TEMPERATURE,
    )
    txt = sanitize_intro(extract_text(resp))

    attempts = 0
    while not _intro_meets_shape(txt, INTRO_MIN_PARAS, INTRO_MIN_WORDS, INTRO_TARGET_MAX) and attempts < INTRO_RETRY_LIMIT:
        attempts += 1
        revision = (
            "REVISION REQUEST:\n"
            f"- Must contain EXACTLY 2 paragraphs.\n"
            f"- Each paragraph must be at least {INTRO_MIN_WORDS} words.\n"
            "- Include at least four UID markers total, using ONLY allowed_uids.\n"
            "- Do not add headings or bullet points.\n"
            "- Output only the revised multi-paragraph text.\n\n"
            "CURRENT TEXT:\n"
            f"{txt}\n"
        )
        resp2 = call_fmapi(
            ENDPOINT,
            messages=[
                {"role": "system", "content": system_text},
                {"role": "user", "content": content},
                {"role": "user", "content": revision},
            ],
            max_tokens=MAX_TOKENS_INTRO,
            temperature=TEMPERATURE,
        )
        txt = sanitize_intro(extract_text(resp2))

    return txt


def generate_summary(system_text: str, cards: list[dict], uid_index: dict, field_value: str) -> str:
    payload = {
        "meta": {"field_filter": field_value, "fiscal_years": []},
        "counts": {"rows": len(cards), "unique_uids": len(uid_index)},
        "institutes_top": top_participating_ics(cards, k=8),
        "allowed_uids": sorted(uid_index.keys()),
    }

    instr = (
        "- Write the SUMMARY for an NIH Triennial report.\n"
        "- Write EXACTLY 2 paragraphs.\n"
        f"- Each paragraph must be substantive and at least {SUMMARY_MIN_WORDS} words.\n"
        "- Use ONLY the provided payload.\n"
        "- Do not invent fiscal years; if fiscal_years is empty, do not mention an FY range.\n"
        "- Include at least two UID markers overall, formatted as [^<UID>] and placed sentence-terminal, using ONLY allowed_uids.\n"
        "- No bullets; no headings; no URLs/PMIDs/JSON/metadata.\n"
        "- Output pure prose paragraphs only.\n"
        "- Do NOT start paragraphs with 'Research', 'Research in', 'Researchers', or 'Studies'.\n"
        "- Avoid repeating the same first-clause structure across paragraphs.\n"
        "- Do NOT introduce citations with phrases like 'as evident in', 'marked by', 'for instance', 'for example', 'such as', or 'including'.\n"
        "- Place UID markers silently at the END of the sentence; do not write 'and/or' between markers.\n"
        "- If multiple citations support one sentence, stack markers, e.g., .[^UID1][^UID2]\n"
    )

    content = instr + "\n" + json.dumps(payload, ensure_ascii=False)
    resp = call_fmapi(
        ENDPOINT,
        messages=[{"role": "system", "content": system_text}, {"role": "user", "content": content}],
        max_tokens=MAX_TOKENS_SUMMARY,
        temperature=TEMPERATURE,
    )
    txt = sanitize_intro(extract_text(resp))

    attempts = 0
    while not _summary_meets_shape(txt, SUMMARY_MIN_PARAS, SUMMARY_MIN_WORDS, SUMMARY_TARGET_MAX) and attempts < SUMMARY_RETRY_LIMIT:
        attempts += 1
        revision = (
            "REVISION REQUEST:\n"
            "- Must contain EXACTLY 2 paragraphs.\n"
            f"- Each paragraph must be at least {SUMMARY_MIN_WORDS} words.\n"
            "- Include at least two UID markers total, using ONLY allowed_uids.\n"
            "- Do not add headings or bullet points.\n"
            "- Output only the revised two-paragraph text.\n\n"
            "CURRENT TEXT:\n"
            f"{txt}\n"
        )
        resp2 = call_fmapi(
            ENDPOINT,
            messages=[
                {"role": "system", "content": system_text},
                {"role": "user", "content": content},
                {"role": "user", "content": revision},
            ],
            max_tokens=MAX_TOKENS_SUMMARY,
            temperature=TEMPERATURE,
        )
        txt = sanitize_intro(extract_text(resp2))

    return txt

def pick_sections(card: dict) -> list[str]:
    # Activity Type removed from routing signal per client requirement
    text = " ".join([
        card.get("Activity Name", ""),
        card.get("Activity Description", ""),
        card.get("Importance", ""),
        card.get("Collaborating ICOs/Agencies/Orgs", ""),
    ]).lower()

    hits = set()

    def has_any(keys):
        return any(k in text for k in keys)

    if has_any(["image", "imaging", "radiology", "ai", "ml", "deep learning", "pet", "mri", "ct", "midrc"]):
        hits.add("Advanced Imaging & AI Tools")
    if has_any(["combination", "combo", "targeted", "inhibitor", "kinase", "precision", "molecularly targeted", "combo therapy"]):
        hits.add("Combination & Targeted Therapies")
    if has_any(["commons", "repository", "portal", "database", "computational", "cloud", "workflow", "data hub", "registry"]):
        hits.add("Data Commons and Computational Resources")
    if has_any(["environmental", "exposure", "toxic", "pollut", "air", "water", "environment", "occupational"]):
        hits.add("Environmental Health and Cancer")
    if has_any(["epidemiology", "surveillance", "registry", "incidence", "prevalence", "cohort", "population"]):
        hits.add("Epidemiology & Surveillance")
    if has_any(["genetic", "genome", "omics", "transcript", "proteomic", "epigen", "cell", "mechanism", "mutation", "gene"]):
        hits.add("Genetics, Cell Biology, and -Omics")
    if has_any(["immunotherapy", "checkpoint", "t cell", "car-t", "immune", "nk cell", "neoantigen"]):
        hits.add("Immunotherapy")
    if has_any(["nutrition", "diet", "exercise", "symptom", "quality of life", "palliative", "cachexia"]):
        hits.add("Nutrition & Symptom Management")
    if has_any(["prevent", "screen", "risk reduction", "vaccin", "hpv", "self-collection"]):
        hits.add("Preventive Interventions")
    if has_any(["recalcitrant", "hard-to-treat", "glioblastoma", "pancreatic", "rare", "refractory"]):
        hits.add("Recalcitrant & Hard-to-Treat Cancer Research")
    if has_any(["screen", "early detection", "biomarker", "liquid biopsy", "mcde"]):
        hits.add("Screening & Early Detection")
    if has_any(["microenvironment", "stroma", "stromal", "macrophage", "myeloid", "tme", "caf"]):
        hits.add("Tumor Microenvironment & Immunology")

    if not hits:
        hits.add("Genetics, Cell Biology, and -Omics")

    return [s for s in SECTION_ORDER if s in hits]


# -----------------------------
# Section routing with rationale
# -----------------------------
ROUTING_MAX_TOKENS = int(os.environ.get("ROUTING_MAX_TOKENS", "220"))

def _safe_json_loads(txt: str) -> dict:
    try:
        return json.loads(txt)
    except Exception:
        # Try to extract first JSON object from mixed output
        m = re.search(r"\{.*\}", txt or "", flags=re.S)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                return {}
        return {}

def llm_route_uid_to_single_section(system_text: str, uid: str, card: dict, candidate_sections: list[str]) -> tuple[str, str, dict]:
    """
    Returns: (selected_section, rationale, excluded_sections_map)
      - selected_section: one item from SECTION_ORDER
      - rationale: 1–2 sentences explaining why this UID belongs in that section
      - excluded_sections_map: {section_name: reason_for_exclusion}
    If the model fails, falls back to heuristic routing (first candidate).
    """
    candidates = [s for s in candidate_sections if s in SECTION_ORDER]
    if not candidates:
        candidates = list(SECTION_ORDER)

    payload = {
        "uid": uid,
        "activity_name": card.get("Activity Name", "—"),
        "activity_description": card.get("Activity Description", "—"),
        "importance": card.get("Importance", "—"),
        "collaborators": card.get("Collaborating ICOs/Agencies/Orgs", "—"),
        "candidate_sections": candidates,
        "all_sections": list(SECTION_ORDER),
    }

    instr = (
        "You are routing ONE activity (UID) into EXACTLY ONE report section.\n"
        "Return STRICT JSON ONLY (no markdown):\n"
        "{\n"
        "  \"selected_section\": \"<one section title from candidate_sections>\",\n"
        "  \"rationale\": \"<1-2 sentences grounded in the activity text>\",\n"
        "  \"excluded_sections\": { \"<other candidate section>\": \"<why not>\", ... }\n"
        "}\n"
        "Rules:\n"
        "- selected_section MUST be one of candidate_sections.\n"
        "- rationale must cite concrete cues (methods, domain terms, resource types) from the activity text.\n"
        "- excluded_sections must include ONLY the other candidate_sections (not all sections).\n"
        "- If multiple candidates fit, pick the BEST match and explain why the others are secondary.\n"
    )

    try:
        resp = call_fmapi(
            ENDPOINT,
            messages=[
                {"role": "system", "content": system_text},
                {"role": "user", "content": instr + "\n" + json.dumps(payload, ensure_ascii=False)},
            ],
            max_tokens=ROUTING_MAX_TOKENS,
            temperature=0.0,
        )
        txt = extract_text(resp).strip()
        data = _safe_json_loads(txt)
        selected = (data.get("selected_section") or "").strip()
        rationale = (data.get("rationale") or "").strip()
        excluded = data.get("excluded_sections") or {}
        if not isinstance(excluded, dict):
            excluded = {}
        if selected not in candidates:
            selected = candidates[0]
        if not rationale:
            # Provide a grounded fallback rationale (heuristic) if model returns empty
            rationale = "Selected based on dominant topical/method cues in the activity title/description."
        # keep only candidate sections (excluding selected)
        excluded = {k: str(v).strip() for k, v in excluded.items() if k in candidates and k != selected and str(v).strip()}
        return selected, rationale, excluded
    except Exception:
        # Pure heuristic fallback
        selected = candidates[0] if candidates else SECTION_ORDER[0]
        rationale = "Selected by rule-based keyword routing when model routing was unavailable."
        excluded = {k: "Not the primary thematic match under keyword routing." for k in candidates if k != selected}
        return selected, rationale, excluded


def enforce_single_section_per_uid(section_to_uids_raw: dict, uid_routing: dict, *, on_conflict: str = "keep_first") -> tuple[dict, dict]:
    """
    Enforce that each UID belongs to EXACTLY ONE section globally.

    Inputs:
      - section_to_uids_raw: {section: [uids]}
      - uid_routing: {uid: {"selected_section": ..., ...}}

    Behavior:
      - If a UID is found in multiple sections, we either:
          * keep_first: keep the first assignment and drop subsequent duplicates
          * raise: raise ValueError

    Returns:
      - (section_to_uids, uid_routing) normalized so each UID appears in one section only.
    """
    from collections import OrderedDict

    # Preserve SECTION_ORDER ordering
    normalized = {sec: [] for sec in SECTION_ORDER}

    seen = OrderedDict()  # uid -> section (first wins)
    conflicts = {}

    for sec in SECTION_ORDER:
        for uid in (section_to_uids_raw.get(sec) or []):
            if uid not in seen:
                seen[uid] = sec
                normalized[sec].append(uid)
            else:
                conflicts.setdefault(uid, set()).update({seen[uid], sec})
                if on_conflict == "raise":
                    raise ValueError(f"UID '{uid}' assigned to multiple sections: {sorted(conflicts[uid])}")
                # keep_first: silently drop uid from this later section

    # Ensure uid_routing matches the normalized assignment (first wins)
    for uid, sec in seen.items():
        if uid in uid_routing:
            uid_routing[uid]["selected_section"] = sec

    return normalized, uid_routing

def route_all_uids(system_text: str, uid_index: dict) -> tuple[dict, dict]:
    """
    Builds:
      - section_to_uids: {section: [uid,...]} where each UID appears exactly once globally
      - uid_routing: {uid: {selected_section, rationale, excluded_sections{sec:reason}, candidates[]}}
    """
    section_to_uids = {sec: [] for sec in SECTION_ORDER}
    uid_routing = {}

    for uid, card in uid_index.items():
        candidates = pick_sections(card)  # heuristic candidates
        selected, rationale, excluded = llm_route_uid_to_single_section(system_text, uid, card, candidates)
        section_to_uids.setdefault(selected, []).append(uid)
        uid_routing[uid] = {
            "selected_section": selected,
            "rationale": rationale,
            "excluded_sections": excluded,
            "candidates": [c for c in candidates if c in SECTION_ORDER],
        }
    # Hard guard: make it impossible for a UID to exist in multiple sections
    section_to_uids, uid_routing = enforce_single_section_per_uid(section_to_uids, uid_routing, on_conflict="keep_first")

    return section_to_uids, uid_routing
def section_synthesis(system_text: str, section_name: str, uids: list[str], uid_index: dict):
    instr = (
            "- Write one or two cohesive synthesis paragraphs for the section title below.\n"
            "- Use ONLY the provided row facts.\n"
            "- Focus on scientific themes, methods, and collaboration patterns.\n"
            "- Do NOT repeat individual activity descriptions verbatim.\n"
            "- Do NOT mention Activity Type or Importance labels.\n"
            "- Optionally include ONE UID marker at the end if a concrete example strengthens the synthesis.\n"
            "- No bullets, no headings, no URLs, no metadata.\n"
            "\n"
            "OPENING VARIATION RULES (STRICT):\n"
            "- Do NOT start any paragraph with 'Research', 'Research in', 'Research on', 'Studies', or 'Researchers'.\n"
            "- Do NOT use the template 'Research in <section> has...'.\n"
            "- Start each paragraph with one of these instead:\n"
            "  (a) A method/approach\n"
            "  (b) A scientific advance/result\n"
            "  (c) A resource/infrastructure\n"
            "- Ensure the first 8 words of each section synthesis paragraph are structurally distinct.\n"
        )

    rows = []
    for u in uids[:6]:
        c = uid_index[u]
        rows.append({
            "UID": u,
            "Submitting ICO": c.get("Submitting ICO", "—"),
            "Lead ICO": c.get("Lead ICO", "—"),
            "Activity Name": c.get("Activity Name", "—"),
            "Activity Description": (c.get("Activity Description", "—")[:400]),
        })
    payload = {"section": section_name, "rows": rows}

    resp = call_fmapi(
        ENDPOINT,
        messages=[
            {"role": "system", "content": system_text},
            {"role": "user", "content": instr + "\n" + json.dumps(payload, ensure_ascii=False)},
        ],
        max_tokens=MAX_TOKENS_SYN,
        temperature=TEMPERATURE,
    )
    txt = extract_text(resp)
    txt = re.sub(r"^\s{0,3}#{1,6}\s+.*$", "", txt, flags=re.MULTILINE).strip()
    return txt


def assemble_markdown(
    summary_text: str,
    intro_text: str,
    section_order: list[str],
    section_to_uids: dict,
    section_syn: dict,
    uid_to_paragraph: dict,
    cards: list | None = None
) -> str:
    md_parts: list[str] = []

    intro_clean = hard_clean_generated_text((intro_text or "").strip())
    intro_clean = normalize_uid_marker_placement(intro_clean)

    summary_clean = hard_clean_generated_text((summary_text or "").strip())
    summary_clean = normalize_uid_marker_placement(summary_clean)

    uid_to_entry = build_uid_entry_map(section_order, section_to_uids)
    intro_clean = nih_style_citation_phrasing(intro_clean)
    summary_clean = nih_style_citation_phrasing(summary_clean)

    intro_clean = strip_raw_uid_tokens(intro_clean)
    summary_clean = strip_raw_uid_tokens(summary_clean)

    intro_clean = ensure_period_after_citations(intro_clean)
    summary_clean = ensure_period_after_citations(summary_clean)

    intro_clean = postprocess_narrative(intro_clean)
    summary_clean = postprocess_narrative(summary_clean)

    intro_clean = enforce_intro_summary_rules(intro_clean)
    summary_clean = enforce_intro_summary_rules(summary_clean)

    intro_clean = finalize_multparagraph_text(intro_clean)
    summary_clean = finalize_multparagraph_text(summary_clean)

    md_parts.append("## Introduction\n")
    if intro_clean:
        md_parts.append(intro_clean + "\n")

    md_parts.append("\n## Summary\n")
    if summary_clean:
        md_parts.append(summary_clean + "\n")

    for sec in section_order:
        uids = section_to_uids.get(sec, []) or []
        if not uids:
            continue

        md_parts.append(f"\n## {sec}\n")
        md_parts.append(", ".join(uids) + "\n")

        for uid in uids:
            entry_n = uid_to_entry.get(uid)
            entry_label = f"Entry {entry_n}" if entry_n is not None else "Entry"

            md_parts.append(f"\nUID {uid}\n")

            para = (uid_to_paragraph.get(uid) or "").strip()
            if (not para) or re.match(r"^See\s+Entry\s+\d+\b", para, flags=re.IGNORECASE):
                card = None
                if isinstance(cards, dict):
                    card = cards.get(uid)
                elif isinstance(cards, list):
                    for _c in cards:
                        if (_c.get("Unique ID") or "").strip() == uid:
                            card = _c
                            break

                para = fallback_paragraph_from_card(card) if card else ""
                if not para:
                    continue

            para = hard_clean_generated_text(para)
            para = re.sub(r"\s+", " ", para).strip()
            para = re.sub(r"\s+([\.,;:!?])", r"\1", para)
            para = strip_raw_uid_tokens(para)
            para = postprocess_narrative(para)
            # Client requirement: keep Entry numbers, but never as citation-like text at paragraph end
            para = re.sub(r"\(\s*Entry\s+\d+\s*\)", "", para, flags=re.IGNORECASE)
            para = re.sub(r"\bEntry\s+\d+\b", "", para, flags=re.IGNORECASE)
            para = re.sub(r"\s{2,}", " ", para).strip()
            para = re.sub(r"\s+([,.;:!?])", r"\1", para)


            para = re.sub(r"\[\^\s*[A-Za-z0-9._-]+\s*\]\s*$", "", para).strip()

            if para.endswith((";", ":")):
                para = para[:-1].rstrip()
            if not para.endswith("."):
                para += "."

            para = apply_uid_markers_from_cite_tokens(para, uid)

            md_parts.append(f"\n{para}\n")


    md_text_so_far = "\n".join(md_parts).strip()
    acr_body = build_acronyms_section(
        md_text_so_far,
        cards=cards,
        system_text=system_text_base if "system_text_base" in globals() else None
    ).strip()

    md_parts.append("\n## Acronyms\n")
    if acr_body:
        md_parts.append(acr_body + "\n")
    else:
        md_parts.append("None identified.\n")

    md_parts.append("\n## References\n\n")

    return "\n".join(md_parts).strip() + "\n"


def ensure_pandoc() -> str:
    """
    Resolve pandoc WITHOUT any runtime downloads.

    Resolution order:
      1) pandoc on PATH
      2) pypandoc-binary bundled pandoc (installed at build time via pip)
      3) $PANDOC_PATH (only if it points to a real runtime filesystem path)
      4) /tmp/pandoc/bin/pandoc (legacy if your image bakes it there)
    """
    import os
    import shutil
    from pathlib import Path

    # 1) PATH
    existing = shutil.which("pandoc")
    if existing:
        return existing

    # 2) pypandoc-binary (bundled pandoc, no runtime download)
    try:
        import pypandoc  # type: ignore
        p = pypandoc.get_pandoc_path()
        if p and Path(p).exists():
            return str(Path(p).resolve())
    except Exception:
        pass

    # 3) Explicit override
    env_path = os.environ.get("PANDOC_PATH", "").strip()
    if env_path:
        p = Path(env_path)
        if p.exists() and p.is_file():
            return str(p.resolve())

    # 4) Legacy baked path
    legacy = Path("/tmp/pandoc/bin/pandoc")
    if legacy.exists() and legacy.is_file():
        return str(legacy)

    raise RuntimeError(
        "Pandoc is required to export DOCX, but it was not found.\n\n"
        "Fix (recommended for Databricks Apps): add pypandoc + pypandoc-binary to requirements.txt "
        "so pandoc is bundled at build time.\n"
    )

def export_docx(md_text: str, out_dir: str, field_value: str, lua_pagebreak_enabled: bool, square_brackets_enabled: bool) -> str:
    pandoc = ensure_pandoc()

    md_text = replace_percent_sign(md_text)
    if "%" in md_text:
        md_text = re.sub(r"(\d(?:\d|\.\d+)?)\s*%", r"\1 percent", md_text)
        md_text = md_text.replace("%", " percent")

    md_text = re.sub(r"^\s*See Entry\s+\d+\s+for\s+the\s+full\s+activity\s+narrative\.?\s*$", "", md_text, flags=re.IGNORECASE | re.MULTILINE)

    field_part = _safe_filename(field_value)
    docx_path = str(Path(out_dir) / f"Triennial_Data_{field_part}.docx")
    md_path = str(Path(out_dir) / "report.md")
    Path(md_path).write_text(md_text, encoding="utf-8")

    # env = os.environ.copy()
    # env["PATH"] = f"/tmp/pandoc/bin:{env.get('PATH', '')}"
    env = os.environ.copy()

    cmd = [
        pandoc, md_path,
        "-o", docx_path,
        "--from", "markdown+footnotes+autolink_bare_uris",
        "--to", "docx",
        "--wrap=none",
        "--standalone",
    ]

    if Path(REFERENCE_DOCX_LOCAL).exists():
        cmd += ["--reference-doc", REFERENCE_DOCX_LOCAL]

    if lua_pagebreak_enabled and Path(LUA_FILTER_LOCAL).exists():
        cmd += ["--lua-filter", LUA_FILTER_LOCAL]


    if square_brackets_enabled and Path(SQUARE_FILTER_LOCAL).exists():
        cmd += ["--lua-filter", SQUARE_FILTER_LOCAL]

    p = subprocess.run(cmd, env=env, capture_output=True, text=True)
    if p.returncode != 0:
        raise RuntimeError(
            "Pandoc failed.\n"
            f"Return code: {p.returncode}\n\n"
            "STDOUT:\n" + (p.stdout or "(empty)") + "\n\n"
            "STDERR:\n" + (p.stderr or "(empty)") + "\n"
        )

    return docx_path


def publish_for_download(local_docx_path: str, out_dbfs_dir: str) -> str:
    dbfs_mkdirs(out_dbfs_dir)
    filename = Path(local_docx_path).name
    target = f"{out_dbfs_dir}/{filename}"
    data = Path(local_docx_path).read_bytes()
    dbfs_put_large(target, data, overwrite=True)

    rel = target.replace("dbfs:/FileStore/", "")
    return f"/files/{rel}"


def maybe_copy_to_volume(local_docx_path: str, volume_dir: str):
    vol = Path(volume_dir)
    if not vol.exists():
        raise RuntimeError(f"Volume directory does not exist in this App container: {volume_dir}")
    if not vol.is_dir():
        raise RuntimeError(f"Volume path is not a directory: {volume_dir}")
    target = vol / Path(local_docx_path).name
    shutil.copyfile(local_docx_path, str(target))


def write_style_prompt_to_dbfs(style_prompt_dbfs_path: str, system_text: str):
    host = get_workspace_host()
    url = f"{host}/api/2.0/dbfs/put"

    payload_obj = {"content": system_text}
    payload_bytes = json.dumps(payload_obj, ensure_ascii=False, indent=2).encode("utf-8")

    r = requests.post(
        url,
        headers=auth_headers(),
        json={
            "path": dbfs_norm(style_prompt_dbfs_path),
            "contents": base64.b64encode(payload_bytes).decode("utf-8"),
            "overwrite": True,
        },
        timeout=30,
    )
    r.raise_for_status()


# =============================
# 9) UI: Inputs + Load button
# =============================
with st.expander("Inputs", expanded=True):
    c1, c2 = st.columns(2)

    with c1:
        excel_path = st.text_input("Excel (app/local path)", value=DEFAULT_EXCEL_PATH)
        style_prompt_path = st.text_input("Style prompt JSON (app/local path)", value=DEFAULT_STYLE_PROMPT_PATH)
        reference_docx_path = st.text_input("Reference DOCX (app/local path)", value=DEFAULT_REFERENCE_DOCX_PATH)
        lua_filter_path = st.text_input("Lua pagebreak filter (app/local path, optional)", value=DEFAULT_LUA_FILTER_PATH)
        square_filter_path = st.text_input("Lua square-bracket filter (app/local path, optional)", value=DEFAULT_SQUARE_FILTER_PATH)

    with c2:
        # DBFS publish is optional; Streamlit's download button works without DBFS
        publish_to_dbfs = st.checkbox("Publish a copy to DBFS FileStore (optional)", value=False)
        working_out_dbfs = st.text_input("Working output dir (DBFS path)", value=DEFAULT_WORKING_OUT_DBFS, disabled=not publish_to_dbfs)

        volume_out_dir = st.text_input("Final output volume dir (optional)", value=DEFAULT_VOLUME_OUT_DIR)
        copy_to_volume = st.checkbox("Also copy final DOCX into Volume directory", value=False)

    use_lua_filter = True
    use_square_bracket_filter = False
    st.caption("Tip: Put the input files next to app.py (or under ./assets/) and keep these as relative names (e.g., 'style_prompt.json').")
    load_inputs = st.button("Load Inputs")


if "inputs_loaded" not in st.session_state:
    st.session_state.inputs_loaded = False

if load_inputs or not st.session_state.inputs_loaded:
    with st.spinner("Loading inputs (staging local/app files)…"):
        stage_assets_or_stop(
            excel_path=excel_path.strip(),
            style_prompt_path=style_prompt_path.strip(),
            reference_docx_path=reference_docx_path.strip(),
            lua_filter_path=lua_filter_path.strip(),
            square_filter_path=square_filter_path.strip(),
        )
        try:
            load_excel.clear()
            load_system_text.clear()
        except Exception:
            pass
        st.session_state.inputs_loaded = True

system_text_base = load_system_text(STYLE_PROMPT_LOCAL)
df = load_excel(EXCEL_LOCAL)

field_col = resolve_any_column(df, ["Field", "Field Type", "FieldType", "field", "field_type"], fallback="Field")

# =============================
# 9.1) Field dropdown (Activity Type filter removed)
# =============================
field_values = dropdown_values(df, field_col)

# =============================
# 9.1) Field dropdown (constrained width)
# =============================
field_values = dropdown_values(df, field_col)

_field_left, _field_gutter = st.columns([0.72, 0.28])  # left column + right gutter
with _field_left:
    selected_field = st.selectbox(
        "Field",
        field_values,
        index=field_values.index("Cancer") if "Cancer" in field_values else 0,
        key="field_select",
    )




# -----------------------------
# Reset plan/routing state when Field changes
# (prevents UID routing from leaking across fields in Streamlit session_state)
# -----------------------------
if "last_selected_field" not in st.session_state:
    st.session_state["last_selected_field"] = None

if st.session_state["last_selected_field"] != str(selected_field):
    for _k in [
        "plan_ready",
        "plan_text",
        "plan_counts",
        "plan_uids",
        "plan_sections",
        "plan_section_counts",
        "sections_confirmed",
        "uid_routing_preview",
        "included_sections",
    ]:
        if _k in st.session_state:
            del st.session_state[_k]
    st.session_state["last_selected_field"] = str(selected_field)

st.divider()

# =============================
# 9.2) Preflight + confirmation gate
# =============================
with st.expander("Preflight (review preview before generating)", expanded=True):
    style_override = st.text_area(
        "Optional: style / tone / structure requests (applied to this run)",
        value="",
        height=120,
        placeholder="Example: use shorter sentences; avoid jargon; emphasize collaboration and translational impact.",
        key="style_override",
    )

    persist_style_override = st.checkbox(
        "Also update style_prompt.json in DBFS using the override (optional)",
        value=False,
        help="If checked, the effective system prompt for this run will be written back to DBFS style_prompt.json.",
    )

    docx_color_text = st.text_input(
        "Optional: DOCX primary text color (formatting) — e.g., blue or #0000ff",
        value="",
        help="This changes Word formatting (unlike LLM prompting). Examples: blue, #0000ff, 0000ff, rgb(0,0,255).",
    )

    build_plan = st.button("Build Plan", type="secondary")

    if "plan_ready" not in st.session_state:
        st.session_state.plan_ready = False
    if "plan_text" not in st.session_state:
        st.session_state.plan_text = ""
    if "plan_counts" not in st.session_state:
        st.session_state.plan_counts = {}
    if "plan_uids" not in st.session_state:
        st.session_state.plan_uids = []

    if "plan_sections" not in st.session_state:
        st.session_state.plan_sections = []
    if "plan_section_counts" not in st.session_state:
        st.session_state.plan_section_counts = {}
    if "sections_confirmed" not in st.session_state:
        st.session_state.sections_confirmed = False
    if "uid_routing_preview" not in st.session_state:
        st.session_state.uid_routing_preview = {}
    if "included_sections" not in st.session_state:
        st.session_state.included_sections = []

    if build_plan:
        filtered_for_plan = df[
            (df[field_col].astype(str) == str(selected_field))
        ].copy()
        filtered_for_plan = filtered_for_plan.fillna("—")

        if len(filtered_for_plan) == 0:
            st.warning("No rows match the selected filters.")
            st.session_state.plan_ready = False
        else:
            cmap = resolve_map(df.columns)
            cards_for_plan = [make_card(r, cmap) for _, r in filtered_for_plan.iterrows()]

            uid_index_for_plan = {}
            for c in cards_for_plan:
                uid = c.get("Unique ID", "—")
                if uid and uid != "—":
                    uid_index_for_plan[uid] = c

            uids = list(uid_index_for_plan.keys())
            counts = {"filtered_rows": int(len(filtered_for_plan)), "unique_uids": int(len(uids))}

            st.write(
                f"Preview: {counts['filtered_rows']} filtered rows, {counts['unique_uids']} unique UIDs "
                f"for Field='{selected_field}'."
            )
            st.caption("First UIDs (up to 30):")
            st.code(", ".join(uids[:30]) if uids else "(none)")

            eff_system_text = build_effective_system_text(system_text_base, style_override)

            with st.spinner("Generating preview (LLM)…"):
                plan_txt = generate_plan(
                    field_value=str(selected_field),
                    uid_list=uids,
                    counts=counts,
                    style_override=style_override,
                )

            st.subheader("Preview (LLM-generated)")
            st.write(plan_txt)

            # --- Section inclusion + rationale preview (requires user confirmation) ---
            effective_system_text_preview = build_effective_system_text(system_text_base, style_override)
            section_to_uids_preview, uid_routing_preview = route_all_uids(effective_system_text_preview, uid_index_for_plan)

            included_sections = [sec for sec in SECTION_ORDER if section_to_uids_preview.get(sec)]
            section_counts = {sec: len(section_to_uids_preview.get(sec, [])) for sec in included_sections}

            st.subheader("sections that will be included")
            if not included_sections:
                st.info("No sections matched the filtered activities; the report will include Introduction and Summary only.")
            else:
                _sec_rows = [{"Section": sec, "Activities (UIDs)": section_counts[sec]} for sec in included_sections]
                #st.dataframe(pd.DataFrame(_sec_rows), use_container_width=True, hide_index=True)
                _plan_tbl_left, _plan_tbl_gutter = st.columns([0.72, 0.28])
                with _plan_tbl_left:
                    st.dataframe(pd.DataFrame(_sec_rows), use_container_width=True, hide_index=True)

            st.subheader("why each UID is in its section")
            _why_rows = []
            for _uid, _meta in uid_routing_preview.items():
                _why_rows.append({
                    "UID": _uid,
                    "Selected Section": _meta.get("selected_section", "—"),
                    "Rationale (LLM)": _meta.get("rationale", "—"),
                    "Excluded Candidates": ", ".join(sorted((_meta.get("excluded_sections") or {}).keys())) or "—",
                })
            #st.dataframe(pd.DataFrame(_why_rows), use_container_width=True, hide_index=True)
            _plan_why_left, _plan_why_gutter = st.columns([0.72, 0.28])
            with _plan_why_left:
                st.dataframe(pd.DataFrame(_why_rows), use_container_width=True, hide_index=True)

            # --- Excluded candidate sections (per UID): always shown AFTER the UID routing table ---
            # NOTE: This block is inside the main Preflight expander, so it must NOT create another expander.
            st.markdown("**Excluded candidate sections (per UID):**")
            any_excluded = False

            for _uid, _meta in uid_routing_preview.items():
                _ex = _meta.get("excluded_sections") or {}
                if not _ex:
                    continue

                any_excluded = True

                # Header line (no bullet prefix)
                st.markdown(f"**{_uid}** (selected: {_meta.get('selected_section','—')}):")

                # Each excluded section as plain lines (no list syntax)
                for _sec, _why in _ex.items():
                    st.write(f"{_sec}: {_why}")

                st.write("")  # spacing between UIDs

            if not any_excluded:
                st.caption("No excluded candidate sections were recorded for the current UID routing.")
            st.session_state.plan_sections = included_sections
            st.session_state.plan_section_counts = section_counts
            st.session_state.uid_routing_preview = uid_routing_preview
            st.session_state.included_sections = included_sections
            st.session_state.sections_confirmed = False  # reset on every new plan

            st.session_state.plan_text = plan_txt
            st.session_state.plan_counts = counts
            st.session_state.plan_uids = uids
            st.session_state.plan_ready = True

    # -----------------------------
    # Persistent preview + rationales
    # (Render from session_state so it remains visible even after Generate Report is clicked)
    # -----------------------------
    # NOTE: Do not create a nested expander here. Keep everything inside the Preflight expander.
    if st.session_state.get("plan_ready") and (not build_plan):
        st.markdown("---")
        st.subheader("Plan preview and routing rationale")

        _counts = st.session_state.get("plan_counts") or {}
        _uids = st.session_state.get("plan_uids") or []
        _plan_txt = st.session_state.get("plan_text") or ""

        st.write(
            f"Preview: {_counts.get('filtered_rows', 0)} filtered rows, {_counts.get('unique_uids', 0)} unique UIDs "
            f"for Field='{selected_field}'."
        )
        st.caption("First UIDs (up to 30):")
        st.code(", ".join(_uids[:30]) if _uids else "(none)")

        st.subheader("Preview (LLM-generated)")
        st.write(_plan_txt)

        _included_sections = st.session_state.get("included_sections") or st.session_state.get("plan_sections") or []
        _section_counts = st.session_state.get("plan_section_counts") or {}

        st.subheader("sections that will be included")
        if not _included_sections:
            st.info("No sections matched the filtered activities; the report will include Introduction and Summary only.")
        else:
            _sec_rows = [{"Section": sec, "Activities (UIDs)": _section_counts.get(sec, 0)} for sec in _included_sections]
            _plan_tbl_left, _plan_tbl_gutter = st.columns([0.72, 0.28])
            with _plan_tbl_left:
                st.dataframe(pd.DataFrame(_sec_rows), use_container_width=True, hide_index=True)

        _uid_routing_preview = st.session_state.get("uid_routing_preview") or {}
        if _uid_routing_preview:
            st.subheader("why each UID is in its section")
            _why_rows = []
            for _uid, _meta in _uid_routing_preview.items():
                _why_rows.append({
                    "UID": _uid,
                    "Selected Section": _meta.get("selected_section", "—"),
                    "Rationale (LLM)": _meta.get("rationale", "—"),
                    "Excluded Candidates": ", ".join(sorted((_meta.get("excluded_sections") or {}).keys())) or "—",
                })

            _plan_why_left, _plan_why_gutter = st.columns([0.72, 0.28])
            with _plan_why_left:
                st.dataframe(pd.DataFrame(_why_rows), use_container_width=True, hide_index=True)

            st.markdown("**Excluded candidate sections (per UID):**")
            any_excluded = False

            for _uid, _meta in _uid_routing_preview.items():
                _ex = _meta.get("excluded_sections") or {}
                if not _ex:
                    continue

                any_excluded = True
                st.markdown(f"**{_uid}** (selected: {_meta.get('selected_section','—')}):")

                for _sec, _why in _ex.items():
                    st.write(f"{_sec}: {_why}")

                st.write("")

            if not any_excluded:
                st.caption("No excluded candidate sections were recorded for the current UID routing.")
        else:
            st.caption("Build Plan to generate routing rationale.")


    confirm_sections = st.checkbox(
        "Yes — confirm the sections above will be included in the report",
        value=bool(st.session_state.sections_confirmed),
        disabled=not st.session_state.plan_ready,
        help="This confirms the section list derived from the filtered activities and routing rules.",
    )
    st.session_state.sections_confirmed = bool(confirm_sections)

    proceed = st.checkbox(
        "Yes — proceed to generate using this plan",
        value=False,
        disabled=not (st.session_state.plan_ready and st.session_state.sections_confirmed),
    )



st.divider()
generate = st.button("Generate Report", type="primary", disabled=not (st.session_state.plan_ready and st.session_state.sections_confirmed and proceed))

narration_placeholder = st.empty()
st.caption("Live narration updates here (replaces in place).")


def set_narration(stage: str, detail: str, context: dict):
    try:
        line = narrator_line(stage=stage, detail=detail, context=context)
        ts = time.strftime("%H:%M:%S")
        narration_placeholder.info(f"[{ts}] {line}")
    except Exception:
        ts = time.strftime("%H:%M:%S")
        narration_placeholder.info(f"[{ts}] {detail}")


# =============================
# 10) Run pipeline
# =============================
if generate:
    progress = st.progress(0, text="Starting…")
    status = st.empty()

    ctx = {"field": str(selected_field), "counts": {}}
    narration_every_n = int(NARRATE_EVERY_N_DEFAULT)

    effective_system_text = build_effective_system_text(system_text_base, style_override)

    if persist_style_override and style_override.strip():
        # Portability: persist the override only in the *local staged* style prompt for this run.
        # (The app source folder is typically read-only in Databricks Apps, and the client may not use DBFS.)
        try:
            payload_obj = {"content": effective_system_text}
            Path(STYLE_PROMPT_LOCAL).write_text(json.dumps(payload_obj, ensure_ascii=False, indent=2), encoding="utf-8")
            st.success("style_prompt.json updated locally for this run.")
        except Exception as e:
            st.warning(f"Could not update local style_prompt.json (continuing without persisting): {e}")


    def parse_hex_color(s: str) -> Optional[Tuple[int, int, int]]:
        if not s:
            return None
        t = s.strip().lower()

        if t in ("blue", "primary blue"):
            return (0, 0, 255)
        if t == "red":
            return (255, 0, 0)
        if t == "green":
            return (0, 128, 0)
        if t == "black":
            return (0, 0, 0)

        m = re.search(r"rgb\(\s*(\d{1,3})\s*,\s*(\d{1,3})\s*,\s*(\d{1,3})\s*\)", t)
        if m:
            r, g, b = [int(x) for x in m.groups()]
            if all(0 <= x <= 255 for x in (r, g, b)):
                return (r, g, b)

        t = t.lstrip("#")
        if re.fullmatch(r"[0-9a-f]{6}", t):
            r = int(t[0:2], 16)
            g = int(t[2:4], 16)
            b = int(t[4:6], 16)
            return (r, g, b)

        return None

    docx_primary_rgb = parse_hex_color(docx_color_text)

    try:
        set_narration("start", "Initializing generation pipeline.", ctx)

        status.write("Filtering rows…")
        set_narration("filter", f"Filtering dataset for Field='{selected_field}'.", ctx)

        filtered = df[
            (df[field_col].astype(str) == str(selected_field))
        ].copy()
        filtered = filtered.fillna("—")

        st.write(f"Filtered rows: {len(filtered)}")
        if len(filtered) == 0:
            set_narration("filter", "No matching rows found for the selected filters.", ctx)
            st.warning("No rows match the selected filters.")
            st.stop()

        progress.progress(0.10, text="Preparing cards…")
        set_narration("cards", "Normalizing columns and preparing row cards.", ctx)

        cmap = resolve_map(df.columns)
        cards = [make_card(r, cmap) for _, r in filtered.iterrows()]

        uid_index = {}
        excluded_uids = {}  # {uid: reason} where possible

        for c in cards:
            uid = (c.get("Unique ID", "—") or "—").strip()

            if not uid or uid == "—":
                # Can't list by UID if missing; track count only
                excluded_uids.setdefault("__MISSING_UID__", 0)
                excluded_uids["__MISSING_UID__"] += 1
                continue

            if uid in uid_index:
                # Duplicate UID: enforce rule that each UID is discussed only once globally
                excluded_uids[uid] = "Duplicate UID encountered in filtered data; only one record is kept for reporting."
                continue

            uid_index[uid] = c

        if not uid_index:
            set_narration("cards", "No valid Unique IDs found in filtered data.", ctx)
            st.error("No valid Unique IDs found in the filtered rows.")
            st.stop()

        ctx["counts"] = {"filtered_rows": int(len(filtered)), "unique_uids": int(len(uid_index))}


        # --- Excluded UID reporting (input/data-level exclusions) ---
        # Note: Missing UID rows are counted under __MISSING_UID__ (no specific UID to display).
        if excluded_uids:
            with st.expander("UIDs excluded from reporting (and why)"):
                rows = []
                missing_n = excluded_uids.get("__MISSING_UID__", 0)
                if missing_n:
                    rows.append({"UID": "—", "Reason": f"{missing_n} row(s) missing Unique ID were excluded (cannot be cited or discussed)."})
                for k, v in excluded_uids.items():
                    if k == "__MISSING_UID__":
                        continue
                    rows.append({"UID": k, "Reason": v})
                if rows:
                    st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)


        progress.progress(0.20, text="Generating row paragraphs…")
        set_narration("rows", f"Generating narrative paragraphs for {len(uid_index)} activities.", ctx)

        uid_to_paragraph = {}
        total = len(uid_index)

        for i, (uid, card) in enumerate(uid_index.items(), 1):
            status.write(f"Generating row paragraph {i}/{total}: {uid}")

            if i == 1 or i == total or (i % narration_every_n == 0):
                set_narration("rows", f"Drafting activity narrative {i} of {total} (UID {uid}).", ctx)

            with st.spinner(f"Model is thinking for UID {uid}…"):
                para = generate_row_paragraph(effective_system_text, card)

            uid_to_paragraph[uid] = para

            if SHOW_PARTIAL_OUTPUT:
                st.write(para)

            progress.progress(0.20 + 0.25 * (i / total), text=f"Row paragraphs: {i}/{total}")

        progress.progress(0.50, text="Routing to sections…")
        set_narration("sections", "Routing activities into report sections (one UID per section) and generating rationales.", ctx)

        section_to_uids, uid_routing = route_all_uids(effective_system_text, uid_index)

        # Persist routing metadata for UI download/review
        st.session_state["uid_routing"] = uid_routing

        evidence_brief = build_portfolio_evidence_brief(filtered, uid_index, section_to_uids)

        progress.progress(0.58, text="Generating Summary…")
        set_narration("summary", "Generating the executive Summary section.", ctx)
        with st.spinner("Model is thinking for the Summary…"):
            summary_text = generate_summary(effective_system_text, cards, uid_index, selected_field)

        progress.progress(0.65, text="Generating Introduction…")
        set_narration("intro", "Generating the multi-paragraph Introduction.", ctx)
        with st.spinner("Model is thinking for the Introduction…"):
            intro_text = generate_intro(effective_system_text, cards, uid_index, selected_field)

        progress.progress(0.78, text="Generating section syntheses…")
        set_narration("synth", "Generating section-level syntheses.", ctx)

        section_syn = {}
        secs_with_uids = [sec for sec in SECTION_ORDER if section_to_uids.get(sec)]
        denom = max(1, len(secs_with_uids))

        done = 0
        for sec in SECTION_ORDER:
            uids = section_to_uids.get(sec, [])
            if not uids:
                continue

            done += 1
            set_narration("synth", f"Writing synthesis for section {done} of {denom}: {sec}.", ctx)
            with st.spinner(f"Model is thinking for section: {sec}…"):
                section_syn[sec] = section_synthesis(effective_system_text, sec, uids, uid_index)

            progress.progress(0.78 + 0.10 * (done / denom), text=f"Section syntheses: {done}/{denom}")

        progress.progress(0.90, text="Assembling markdown…")
        set_narration("assemble", "Assembling report markdown and integrating generated components.", ctx)

        md = assemble_markdown(summary_text, intro_text, SECTION_ORDER, section_to_uids, section_syn, uid_to_paragraph, cards=cards)

        set_narration("footnotes", "Converting UID markers into numeric Word footnotes.", ctx)
        md_numeric, footnote_block, references_md = build_footnotes_from_uid_markers(md, uid_index)
        md_with_refs = _inject_references_section(md_numeric, references_md)
        md = md_with_refs.rstrip() + "\n\n" + footnote_block.strip() + "\n"

        md = enforce_ico_acronyms(md)
        md = revert_ci_expansions(md)
        md = replace_percent_sign(md)

        progress.progress(0.93, text="Exporting DOCX…")
        set_narration("docx", "Exporting the report to DOCX via Pandoc.", ctx)

        with st.spinner("Building DOCX (Pandoc)…"):
            docx_path = export_docx(
                md,
                LOCAL_OUT_DIR,
                selected_field,
                lua_pagebreak_enabled=True,
                square_brackets_enabled=False,
            )

        if docx_primary_rgb is not None and DOCX_AVAILABLE:
            set_narration("format", "Applying DOCX primary text color formatting.", ctx)
            with st.spinner("Applying DOCX text color…"):
                apply_primary_color_to_docx(docx_path, docx_primary_rgb)
        elif docx_primary_rgb is not None and not DOCX_AVAILABLE:
            st.warning(
                "DOCX color override requested, but python-docx is not available in this environment. "
                 "The report was generated successfully without color formatting."
            )

        if copy_to_volume:
            set_narration("volume", "Copying the DOCX into the configured Volume directory.", ctx)
            with st.spinner("Copying to Volume…"):
                maybe_copy_to_volume(docx_path, volume_out_dir)

        progress.progress(0.97, text="Publishing for download…")
        set_narration("publish", "Publishing the DOCX to DBFS FileStore for download.", ctx)

        with st.spinner("Publishing file to DBFS…"):
            url = publish_for_download(docx_path, working_out_dbfs.strip())

        progress.progress(1.0, text="Done.")
        status.success("Report generated successfully.")
        set_narration("done", "Generation complete. The report is ready for download.", ctx)

        st.markdown("### Download")
        docx_bytes = Path(docx_path).read_bytes()
        st.download_button(
            label="Download DOCX",
            data=docx_bytes,
            file_name=Path(docx_path).name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        if url:
            full_url = f"{get_workspace_host()}{url}"
            st.caption(url)
        else:
            st.caption("DBFS publish disabled (download using the button above).")

        # NOTE: UID routing rationale is intentionally shown during **Build Plan** (preflight),
        # not after generation. This keeps the post-generation UI focused on download + diagnostics.
    
    except Exception as e:
        progress.progress(1.0, text="Failed.")
        set_narration("error", f"Generation failed: {e}", ctx)
        st.error(f"Generation failed: {e}")
        st.stop()
